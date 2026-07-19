"""SurgiNote backend - FastAPI app for surgical resident personal assistant."""
import os
import io
import uuid
import asyncio
import logging
from pathlib import Path
from datetime import datetime, timezone, date, timedelta
from typing import Optional

import jwt
import bcrypt
from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field, EmailStr, ConfigDict
from openai import AsyncOpenAI
from docx import Document as DocxDocument

openai_client = AsyncOpenAI(api_key=OPENAI_API_KEY)

from style_examples import (
    hospitalization_note_style_block,
    uti_note_style_block,
    whatsapp_style_block,
)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

mongo_url = os.environ["MONGO_URL"]
mongo_client = AsyncIOMotorClient(mongo_url)

db = mongo_client["surginote"]
db = client[os.environ["DB_NAME"]]

JWT_SECRET = os.environ["JWT_SECRET"]
JWT_ALGORITHM = "HS256"
JWT_EXPIRES_HOURS = 24 * 30

app = FastAPI(title="SurgiNote API")
api = APIRouter(prefix="/api")
security = HTTPBearer()

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger("surginote")


def now_utc() -> datetime:
    return datetime.now(timezone.utc)


def today_iso() -> str:
    return now_utc().date().isoformat()


class UserRegister(BaseModel):
    email: EmailStr
    password: str
    name: str


class UserLogin(BaseModel):
    email: EmailStr
    password: str


class AuthResponse(BaseModel):
    token: str
    user: dict


class Patient(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    name: str
    age: Optional[int] = None
    sex: Optional[str] = None
    bed: Optional[str] = None
    floor: Optional[str] = None
    service: Optional[str] = None
    attending_physician: Optional[str] = None
    dx_short: Optional[str] = None
    dx_full: Optional[str] = None
    surgery_date: Optional[str] = None
    surgery_procedure: Optional[str] = None
    surgery_findings: Optional[str] = None
    medical_history: Optional[str] = None
    allergies: Optional[str] = None
    important_medications: Optional[str] = None
    consultants: Optional[str] = None
    oncology_treatment: Optional[str] = None
    oncology_status: Optional[str] = None
    unit_classification: Optional[str] = None
    admission_date: Optional[str] = None
    admission_note_text: Optional[str] = None
    is_surgical: bool = True
    is_pending_discharge: bool = False
    censo_order: Optional[int] = None
    active: bool = True
    is_new_admission: bool = False
    discharged_at: Optional[str] = None
    created_at: str = Field(default_factory=lambda: now_utc().isoformat())
    updated_at: str = Field(default_factory=lambda: now_utc().isoformat())


class PatientCreate(BaseModel):
    name: str
    age: Optional[int] = None
    sex: Optional[str] = None
    bed: Optional[str] = None
    floor: Optional[str] = None
    service: Optional[str] = None
    attending_physician: Optional[str] = None
    dx_short: Optional[str] = None
    dx_full: Optional[str] = None
    surgery_date: Optional[str] = None
    surgery_procedure: Optional[str] = None
    surgery_findings: Optional[str] = None
    medical_history: Optional[str] = None
    allergies: Optional[str] = None
    important_medications: Optional[str] = None
    consultants: Optional[str] = None
    oncology_treatment: Optional[str] = None
    oncology_status: Optional[str] = None
    unit_classification: Optional[str] = None
    admission_date: Optional[str] = None
    is_surgical: Optional[bool] = None
    is_pending_discharge: Optional[bool] = None


class PatientUpdate(BaseModel):
    name: Optional[str] = None
    age: Optional[int] = None
    sex: Optional[str] = None
    bed: Optional[str] = None
    floor: Optional[str] = None
    service: Optional[str] = None
    attending_physician: Optional[str] = None
    dx_short: Optional[str] = None
    dx_full: Optional[str] = None
    surgery_date: Optional[str] = None
    surgery_procedure: Optional[str] = None
    surgery_findings: Optional[str] = None
    medical_history: Optional[str] = None
    allergies: Optional[str] = None
    important_medications: Optional[str] = None
    consultants: Optional[str] = None
    oncology_treatment: Optional[str] = None
    oncology_status: Optional[str] = None
    unit_classification: Optional[str] = None
    admission_date: Optional[str] = None
    active: Optional[bool] = None
    is_new_admission: Optional[bool] = None
    is_surgical: Optional[bool] = None
    is_pending_discharge: Optional[bool] = None


class DailyEntry(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    patient_id: str
    user_id: str
    date: str
    dictation: str = ""
    labs: str = ""
    studies: str = ""
    procedures: str = ""
    cultures: str = ""
    vital_signs: str = ""
    indications_changes: str = ""
    events: str = ""
    ai_pase_summary: Optional[str] = None
    ai_lab_comparison: Optional[str] = None
    ai_evolution_note: Optional[str] = None
    ai_whatsapp: Optional[str] = None
    saved_to_pase: bool = False
    created_at: str = Field(default_factory=lambda: now_utc().isoformat())
    updated_at: str = Field(default_factory=lambda: now_utc().isoformat())


class DailyEntryUpsert(BaseModel):
    date: Optional[str] = None
    dictation: Optional[str] = None
    labs: Optional[str] = None
    studies: Optional[str] = None
    procedures: Optional[str] = None
    cultures: Optional[str] = None
    vital_signs: Optional[str] = None
    indications_changes: Optional[str] = None
    events: Optional[str] = None


def hash_pw(pw: str) -> str:
    return bcrypt.hashpw(pw.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def check_pw(pw: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(pw.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False


def make_token(user_id: str, email: str) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "exp": now_utc() + timedelta(hours=JWT_EXPIRES_HOURS),
        "iat": now_utc(),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)


async def get_user(creds: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    try:
        payload = jwt.decode(creds.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload["sub"]
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Token inválido")
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=401, detail="Usuario no encontrado")
    return user


def days_between(iso_date: Optional[str]) -> Optional[int]:
    if not iso_date:
        return None
    try:
        d = date.fromisoformat(iso_date)
        return (now_utc().date() - d).days
    except Exception:
        return None


def enrich_patient(p: dict) -> dict:
    p["days_admission"] = days_between(p.get("admission_date"))
    p["days_postop"] = days_between(p.get("surgery_date"))
    return p


@api.post("/auth/register", response_model=AuthResponse)
async def register(body: UserRegister):
    existing = await db.users.find_one({"email": body.email.lower()})
    if existing:
        raise HTTPException(status_code=400, detail="Email ya registrado")
    user_id = str(uuid.uuid4())
    doc = {
        "id": user_id,
        "email": body.email.lower(),
        "password": hash_pw(body.password),
        "name": body.name,
        "created_at": now_utc().isoformat(),
    }
    await db.users.insert_one(doc)
    token = make_token(user_id, body.email.lower())
    return AuthResponse(token=token, user={"id": user_id, "email": body.email.lower(), "name": body.name})


@api.post("/auth/login", response_model=AuthResponse)
async def login(body: UserLogin):
    user = await db.users.find_one({"email": body.email.lower()})
    if not user or not check_pw(body.password, user["password"]):
        raise HTTPException(status_code=401, detail="Credenciales inválidas")
    token = make_token(user["id"], user["email"])
    return AuthResponse(token=token, user={"id": user["id"], "email": user["email"], "name": user["name"]})


# ---------------- Training Examples (Configuración → Entrenamiento de IA) ----------------
class TrainingExamples(BaseModel):
    hospital_notes_examples: str = ""
    uti_notes_examples: str = ""
    whatsapp_examples: str = ""


@api.get("/training-examples", response_model=TrainingExamples)
async def get_training_examples(user: dict = Depends(get_user)):
    doc = await db.training_examples.find_one({"user_id": user["id"]}, {"_id": 0}) or {}
    return TrainingExamples(
        hospital_notes_examples=doc.get("hospital_notes_examples", ""),
        uti_notes_examples=doc.get("uti_notes_examples", ""),
        whatsapp_examples=doc.get("whatsapp_examples", ""),
    )


@api.put("/training-examples", response_model=TrainingExamples)
async def update_training_examples(body: TrainingExamples, user: dict = Depends(get_user)):
    await db.training_examples.update_one(
        {"user_id": user["id"]},
        {"$set": {
            "user_id": user["id"],
            "hospital_notes_examples": (body.hospital_notes_examples or "").strip(),
            "uti_notes_examples": (body.uti_notes_examples or "").strip(),
            "whatsapp_examples": (body.whatsapp_examples or "").strip(),
            "updated_at": now_utc().isoformat(),
        }},
        upsert=True,
    )
    return body


async def _get_training_examples(user_id: str) -> dict:
    """Fetch user's training examples. Returns {} if none configured."""
    doc = await db.training_examples.find_one({"user_id": user_id}, {"_id": 0})
    return doc or {}


@api.get("/auth/me")
async def me(user: dict = Depends(get_user)):
    return user


@api.get("/patients/discharged")
async def list_discharged(user: dict = Depends(get_user)):
    cur = db.patients.find({"user_id": user["id"], "active": False}, {"_id": 0})
    patients = await cur.to_list(1000)
    for p in patients:
        enrich_patient(p)
    patients.sort(key=lambda x: x.get("discharged_at") or "", reverse=True)
    return patients


import re as _re

_BED_SUITE_RE = _re.compile(r"^S(\d)(\d+)$", _re.IGNORECASE)
_BED_ROOM_RE = _re.compile(r"^(\d)(\d{2,})$")


def _bed_sort_key(bed: Optional[str]):
    """Physical hospital order:
    - Descending floor (piso 8, 7, 6, ...)
    - Within each floor: SUITES first (S81, S82, S83), then normal rooms (800, 801, 802)
    - Patients without a valid bed number go last.
    Bed format examples:
      'S81' -> floor=8, suite=True, num=1
      '800' -> floor=8, suite=False, num=0
      'Dajer' or other -> unranked
    """
    if not bed:
        return (999, 9, 999999, "")
    b = str(bed).strip().upper()
    m = _BED_SUITE_RE.match(b)
    if m:
        floor = int(m.group(1))
        num = int(m.group(2))
        return (-floor, 0, -num, b)  # suite = 0 (first within floor), num DESC
    m = _BED_ROOM_RE.match(b)
    if m:
        floor = int(m.group(1))
        num = int(m.group(2))
        return (-floor, 1, -num, b)  # normal room = 1 (after suites), num DESC
    return (999, 9, 999999, b)


@api.get("/patients")
async def list_patients(user: dict = Depends(get_user)):
    cur = db.patients.find({"user_id": user["id"], "active": True}, {"_id": 0})
    patients = await cur.to_list(500)
    for p in patients:
        enrich_patient(p)
    # PHYSICAL hospital order: suites first per floor (descending floors).
    patients.sort(key=lambda p: _bed_sort_key(p.get("bed")))
    return patients


@api.post("/patients", response_model=Patient)
async def create_patient(body: PatientCreate, user: dict = Depends(get_user)):
    payload = {k: v for k, v in body.model_dump().items() if v is not None}
    patient = Patient(user_id=user["id"], **payload)
    await db.patients.insert_one(patient.model_dump())
    return patient


@api.get("/patients/{patient_id}")
async def get_patient(patient_id: str, user: dict = Depends(get_user)):
    p = await db.patients.find_one({"id": patient_id, "user_id": user["id"]}, {"_id": 0})
    if not p:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    return enrich_patient(p)


@api.patch("/patients/{patient_id}")
async def update_patient(patient_id: str, body: PatientUpdate, user: dict = Depends(get_user)):
    update = {k: v for k, v in body.model_dump().items() if v is not None}
    update["updated_at"] = now_utc().isoformat()
    result = await db.patients.update_one(
        {"id": patient_id, "user_id": user["id"]}, {"$set": update}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    p = await db.patients.find_one({"id": patient_id}, {"_id": 0})
    return enrich_patient(p)


@api.delete("/patients/{patient_id}")
async def delete_patient(patient_id: str, user: dict = Depends(get_user)):
    result = await db.patients.update_one(
        {"id": patient_id, "user_id": user["id"]},
        {"$set": {"active": False, "discharged_at": today_iso(), "updated_at": now_utc().isoformat()}},
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    return {"ok": True}


@api.get("/patients/{patient_id}/entries")
async def list_entries(patient_id: str, user: dict = Depends(get_user)):
    cur = db.daily_entries.find({"patient_id": patient_id, "user_id": user["id"]}, {"_id": 0})
    entries = await cur.to_list(365)
    entries.sort(key=lambda x: x.get("date", ""), reverse=True)
    return entries


@api.get("/patients/{patient_id}/entries/today")
async def get_today_entry(patient_id: str, user: dict = Depends(get_user)):
    today = today_iso()
    entry = await db.daily_entries.find_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": today}, {"_id": 0}
    )
    if not entry:
        e = DailyEntry(patient_id=patient_id, user_id=user["id"], date=today)
        await db.daily_entries.insert_one(e.model_dump())
        return e.model_dump()
    return entry


@api.patch("/patients/{patient_id}/entries/today")
async def update_today_entry(patient_id: str, body: DailyEntryUpsert, user: dict = Depends(get_user)):
    target_date = body.date or today_iso()
    update = {k: v for k, v in body.model_dump().items() if v is not None and k != "date"}
    update["updated_at"] = now_utc().isoformat()

    existing = await db.daily_entries.find_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": target_date}
    )
    if not existing:
        entry = DailyEntry(patient_id=patient_id, user_id=user["id"], date=target_date, **update)
        await db.daily_entries.insert_one(entry.model_dump())
        return entry.model_dump()
    await db.daily_entries.update_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": target_date}, {"$set": update}
    )
    e = await db.daily_entries.find_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": target_date}, {"_id": 0}
    )
    return e


def build_patient_context(patient: dict) -> str:
    days_a = days_between(patient.get("admission_date"))
    days_p = days_between(patient.get("surgery_date"))
    lines = [
        f"Nombre: {patient.get('name', '')}",
        f"Edad: {patient.get('age', 'ND')}   Sexo: {patient.get('sex', 'ND')}",
        f"Cama: {patient.get('bed', 'ND')}   Piso: {patient.get('floor', 'ND')}   Servicio: {patient.get('service', 'ND')}",
        f"Unidad: {patient.get('unit_classification', 'Piso')}",
        f"Días de estancia (DEIH): {days_a if days_a is not None else 'ND'}   Días postoperatorios (DPQX): {days_p if days_p is not None else 'ND'}",
        f"Diagnóstico resumido: {patient.get('dx_short', '')}",
        f"Diagnóstico completo: {patient.get('dx_full', '')}",
        f"Procedimiento quirúrgico ({patient.get('surgery_date', 'ND')}): {patient.get('surgery_procedure', '')}",
        f"Hallazgos quirúrgicos: {patient.get('surgery_findings', '')}",
        f"Antecedentes personales patológicos: {patient.get('medical_history', '')}",
        f"Alergias: {patient.get('allergies', 'Ninguna referida')}",
        f"Medicamentos importantes: {patient.get('important_medications', '')}",
        f"Interconsultantes: {patient.get('consultants', '')}",
        f"Tratamiento oncológico: {patient.get('oncology_treatment', '')}",
        f"Estado oncológico: {patient.get('oncology_status', '')}",
    ]
    return "\n".join(lines)


async def llm_generate(system_prompt: str, user_prompt: str) -> str:
    response = await client.chat.completions.create(
        model="gpt-5",
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ]
    )

    return response.choices[0].message.content


class GenerateBody(BaseModel):
    date: Optional[str] = None


def _is_uti_patient(patient: dict) -> bool:
    unit = (patient.get("unit_classification") or "").strip().upper()
    return unit in ("UTI", "UTIM", "TERAPIA", "TERAPIA INTENSIVA", "UCI", "UCIM")


def _style_block(examples: str, kind: str) -> str:
    """Build a style-reference block for the LLM prompt. Never copy verbatim."""
    if not examples or not examples.strip():
        return ""
    return (
        f"\n\nEJEMPLOS DE ESTILO DE {kind} (SOLO REFERENCIA — NUNCA COPIAR NOMBRES, DIAGNÓSTICOS, "
        f"FECHAS NI FRASES COMPLETAS; ÚNICAMENTE APRENDER ESTRUCTURA, VOCABULARIO Y FORMA DE REDACTAR "
        f"COMO UN RESIDENTE R3 DE CIRUGÍA GENERAL):\n"
        f"<<<INICIO_EJEMPLOS>>>\n{examples.strip()}\n<<<FIN_EJEMPLOS>>>\n"
    )


def _extract_dr_lastname(attending: Optional[str]) -> str:
    """Extract 'Apellido' from something like 'Dr. Juan Pérez López' → 'Pérez'."""
    if not attending:
        return "________"
    s = _re.sub(r"^\s*(dra?\.?|dr\.?)\s+", "", attending.strip(), flags=_re.IGNORECASE)
    parts = [w for w in s.split() if w]
    if not parts:
        return "________"
    # Take second word (surname) if pattern is Name Surname, else first
    if len(parts) >= 2:
        return parts[1].upper()
    return parts[0].upper()


@api.post("/patients/{patient_id}/generate/pase")
async def generate_pase(patient_id: str, body: GenerateBody, user: dict = Depends(get_user)):
    patient = await db.patients.find_one({"id": patient_id, "user_id": user["id"]}, {"_id": 0})
    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    target = body.date or today_iso()
    today_entry = await db.daily_entries.find_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": target}, {"_id": 0}
    ) or {"dictation": "", "labs": "", "studies": "", "events": ""}

    prev_cur = db.daily_entries.find(
        {"patient_id": patient_id, "user_id": user["id"], "date": {"$lt": target}, "labs": {"$ne": ""}}, {"_id": 0}
    )
    prev_list = await prev_cur.to_list(30)
    prev_list.sort(key=lambda x: x.get("date", ""), reverse=True)
    prev_entry = prev_list[0] if prev_list else None

    ctx = build_patient_context(patient)
    prev_labs_block = (
        f"\n\nLABORATORIOS DEL DÍA PREVIO ({prev_entry['date']}):\n{prev_entry['labs']}"
        if prev_entry else "\n\n(SIN LABORATORIOS PREVIOS PARA COMPARAR.)"
    )

    dea = days_between(patient.get('admission_date'))
    dpq = days_between(patient.get('surgery_date'))

    # Prompt 5: Pase automático con estructura fija (secuencia obligatoria) + PLAN
    user_prompt = f"""INFORMACIÓN FIJA DEL PACIENTE:
{ctx}

INFORMACIÓN DEL DÍA ({target}):
- DICTADO DEL RESIDENTE: {today_entry.get('dictation', '') or 'SIN DICTADO.'}
- LABORATORIOS: {today_entry.get('labs', '') or 'SIN NUEVOS LABORATORIOS.'}
- ESTUDIOS DE IMAGEN / PROCEDIMIENTOS: {today_entry.get('studies', '') or 'NINGUNO.'}
- CULTIVOS: {today_entry.get('cultures', '') or 'NINGUNO.'}
- PROCEDIMIENTOS: {today_entry.get('procedures', '') or 'NINGUNO.'}
- SIGNOS VITALES: {today_entry.get('vital_signs', '') or 'ND.'}
- EVENTOS Y CAMBIOS DEL DÍA: {today_entry.get('events', '') or 'SIN EVENTOS.'}
- CAMBIOS DE INDICACIONES: {today_entry.get('indications_changes', '') or 'SIN CAMBIOS.'}
{prev_labs_block}

GENERA UN PASE DE VISITA COMPLETO EN UN ÚNICO PÁRRAFO CONTINUO EN MAYÚSCULAS, LISTO PARA LEER EN VOZ ALTA COMO LO HARÍA UN R3 DE CIRUGÍA GENERAL. NO USES BULLETS, LISTAS NI ENCABEZADOS. SIGUE ESTA SECUENCIA OBLIGATORIA:

1. INICIA CON "PACIENTE CON SIGNOS VITALES..." INTERPRETANDO (no copies cifras) — "DENTRO DE PARÁMETROS NORMALES", "CON TENDENCIA A HIPERTENSIÓN", "CON TENDENCIA A HIPOTENSIÓN", "CON TENDENCIA A TAQUICARDIA", "CON TENDENCIA A BRADICARDIA", "CON TENDENCIA A HIPOXIA". Solo agrega una cifra si hay pico máximo/mínimo clínicamente importante (p. ej. "CON TA MÁXIMA DE 170/90 MMHG").
2. OXÍGENO — solo si aplica: "SIN REQUERIMIENTO DE OXÍGENO" o "CON PUNTAS NASALES A 2 LPM" o "CON PUNTAS NASALES DE ALTO FLUJO A 35 LPM FIO2 80%" o "CON VENTILACIÓN MECÁNICA", terminando con "MANTENIENDO ADECUADA SATURACIÓN".
3. DOLOR — "REFIERE ADECUADO CONTROL ANALGÉSICO" O "REFIERE MAL CONTROL ANALGÉSICO, ENA X, POR LO QUE SE INDICÓ..." (si hubo rescate/ajuste, mencionarlo).
4. DIETA — "SE MANTIENE EN AYUNO" o "TOLERANDO ADECUADAMENTE DIETA (LÍQUIDA/BLANDA/ASTRINGENTE/NORMAL)". Si inició VO ayer: "EL DÍA DE AYER SE INICIÓ DIETA...". Conservar abreviaturas BOOST, NPT, NET sin expandir.
5. NÁUSEAS/VÓMITO — "SIN REFERIR NÁUSEAS O VÓMITO" o "PRESENTÓ UN EPISODIO DE VÓMITO, POR LO QUE SE INDICÓ...".
6. DIURESIS — "DIURESIS POR MICCIÓN ESPONTÁNEA PRESENTE" o "DIURESIS POR SONDA FOLEY DE CARACTERÍSTICAS CLARAS/HEMATÚRICAS CUANTIFICADA EN X ML" (mencionar total del cierre de turno). Integrar DKH si aplica.
7. GASES — "CANALIZA GASES" o "PENDIENTE CANALIZACIÓN DE GASES".
8. EVACUACIONES — "EVACUACIONES PRESENTES DE CARACTERÍSTICAS BRISTOL X" / "EVACUACIONES LÍQUIDAS/BLANDAS" / "EVACUACIONES PENDIENTES".
9. DRENAJES/SONDAS (solo si el paciente los tiene) — describirlos con gasto y características.
10. LABORATORIOS — NUNCA copiar valores completos. INTERPRETAR vs previo con este formato exacto: "EL DÍA DE AYER CON DISMINUCIÓN DE LEUCOCITOS DE X A X", "AUMENTO DE HEMOGLOBINA DE X A X", "PERSISTE LEUCOCITOSIS (X MIL)", "PERSISTE HIPOALBUMINEMIA (X)", "MEJORÍA DE FUNCIÓN RENAL", "EMPEORAMIENTO DE FUNCIÓN RENAL", "AUMENTO DE CREATININA DE X A X". Solo cambios relevantes. Omitir normales.
11. ESTUDIOS — si hubo estudio ayer: "EL DÍA DE AYER SE REALIZÓ TAC/USG/RM/RX/ECOCARDIOGRAMA... CON HALLAZGOS DE..." (resumido, nunca reporte completo).
12. CAMBIOS DE MEDICAMENTOS — "SE AGREGÓ OLMESARTÁN POR MAL CONTROL HIPERTENSIVO", "SE INICIÓ INSULINA GLARGINA", "SE SUSPENDIÓ VASOPRESINA", etc.
13. INTERCONSULTAS — "SE INTERCONSULTÓ A MEDICINA INTERNA/NUTRICIÓN/ENDOCRINOLOGÍA/CLÍNICA DE HERIDAS/NEUMOLOGÍA POR..." (si aplica).
14. PROCEDIMIENTOS DE OTROS SERVICIOS — "EL DÍA DE AYER SE LE REALIZÓ DRENAJE ENDOSCÓPICO / CPRE / HEMODIÁLISIS / ANGIOGRAFÍA..." (si aplica).
15. REPORTE HISTOPATOLÓGICO — "SALIÓ EL REPORTE HISTOPATOLÓGICO CON REPORTE DE..." (resumir dx, si aplica).
16. POSTQUIRÚRGICO reciente (si aplica): "EL DÍA DE AYER PASÓ A [PROCEDIMIENTO] CON HALLAZGOS DE [HALLAZGOS]. EN SU POSTQUIRÚRGICO SE HA MANTENIDO CON SIGNOS VITALES..., SE INICIÓ DIETA..., DIURESIS..., DRENAJES...".

INMEDIATAMENTE DESPUÉS DEL PÁRRAFO, EN LÍNEA APARTE, ESCRIBE:

PLAN: CONTINUAR DIETA X, ANALGESIA IV, ANTIBIÓTICOS IV, CUANTIFICACIÓN DE DRENAJES/SONDA, VIGILANCIA DE DATOS DE SANGRADO/ABDOMEN AGUDO/DIFICULTAD RESPIRATORIA, LABORATORIOS DE CONTROL EN AM, VALORAR EGRESO SI APLICA.
(Máximo 2 renglones, separado por comas, sin viñetas, integra SOLO indicaciones vigentes, adáptalo al paciente.)

REGLAS INAMOVIBLES:
- TODO EN MAYÚSCULAS.
- Sin listas, sin viñetas, sin numeración, sin markdown, sin encabezados.
- Solo un párrafo continuo + PLAN.
- Nunca inventes datos. Si algo no está, no lo menciones (no escribas "ND" ni "sin datos aportados").
- Nunca uses frases tipo ChatGPT ("OJO", "LLAMA LA ATENCIÓN", "SE SUGIERE", "SE RECOMIENDA").
- Nunca copies texto del expediente. Reescribe todo con lenguaje de residente."""

    system = (
        "Eres RESIDENTE R3 DE CIRUGÍA GENERAL del HOSPITAL ÁNGELES presentando el pase de "
        "visita en voz alta al equipo. Escribe TODO EN MAYÚSCULAS en un solo párrafo continuo "
        "siguiendo la secuencia obligatoria. Interpretas signos vitales y laboratorios (no "
        "copias cifras crudas). Nunca inventes datos. Nunca uses frases de IA."
    )
    text = await llm_generate(system, user_prompt)

    await db.daily_entries.update_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": target},
        {"$set": {"ai_pase_summary": text, "saved_to_pase": True, "updated_at": now_utc().isoformat()}},
        upsert=True,
    )
    # Al generar el primer pase, quitar el badge "Nuevo ingreso"
    if patient.get("is_new_admission"):
        await db.patients.update_one(
            {"id": patient_id, "user_id": user["id"]},
            {"$set": {"is_new_admission": False, "updated_at": now_utc().isoformat()}},
        )
    return {"summary": text}


@api.post("/patients/{patient_id}/generate/note")
async def generate_note(patient_id: str, body: GenerateBody, user: dict = Depends(get_user)):
    patient = await db.patients.find_one({"id": patient_id, "user_id": user["id"]}, {"_id": 0})
    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    target = body.date or today_iso()
    entry = await db.daily_entries.find_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": target}, {"_id": 0}
    ) or {"dictation": "", "labs": "", "studies": "", "procedures": "", "cultures": "", "vital_signs": "", "events": ""}

    ctx = build_patient_context(patient)
    is_uti = _is_uti_patient(patient)
    training = await _get_training_examples(user["id"])
    user_style_examples = training.get("uti_notes_examples" if is_uti else "hospital_notes_examples", "")
    user_style_block = _style_block(user_style_examples, "NOTAS DEL PROPIO USUARIO")

    # Sistema de few-shot permanentes (Prompt 4 — estándar del servicio de Cirugía Ángeles)
    builtin_style_block = uti_note_style_block() if is_uti else hospitalization_note_style_block()

    plantilla = "UTI/UTIM" if is_uti else "HOSPITALIZACIÓN (PISO)"

    data_block = (
        f"INFORMACIÓN FIJA DEL PACIENTE:\n{ctx}\n\n"
        f"INFORMACIÓN CLÍNICA DEL DÍA ({target}):\n"
        f"- DICTADO DEL RESIDENTE: {entry.get('dictation', '') or 'SIN DICTADO NUEVO.'}\n"
        f"- LABORATORIOS: {entry.get('labs', '') or 'SIN LABORATORIOS.'}\n"
        f"- ESTUDIOS DE IMAGEN: {entry.get('studies', '') or 'SIN ESTUDIOS.'}\n"
        f"- PROCEDIMIENTOS: {entry.get('procedures', '') or 'SIN PROCEDIMIENTOS.'}\n"
        f"- CULTIVOS: {entry.get('cultures', '') or 'SIN CULTIVOS.'}\n"
        f"- SIGNOS VITALES: {entry.get('vital_signs', '') or 'ND.'}\n"
        f"- CAMBIOS DE INDICACIONES: {entry.get('indications_changes', '') or 'SIN CAMBIOS.'}\n"
        f"- EVENTOS Y PENDIENTES: {entry.get('events', '') or 'SIN EVENTOS ADICIONALES.'}\n"
    )

    if is_uti:
        format_rules = (
            "FORMATO OBLIGATORIO (imitando los ejemplos permanentes de UTI/UTIM del servicio):\n"
            "1. Inicia con el encabezado exacto: 'NOTA DE EVOLUCIÓN POR CIRUGÍA GENERAL EN UTI' "
            "(o EN UTIM si aplica), en la línea 1.\n"
            "2. Segundo párrafo — INTRODUCCIÓN EXACTA: 'SE TRATA DE PACIENTE MASCULINO/FEMENINO DE "
            "XX AÑOS, NOMBRE COMPLETO, CON DIAGNÓSTICO DE __________ POR LO QUE EL DÍA __________ "
            "SE LE REALIZÓ __________ CON HALLAZGOS DE __________.' NUNCA separes cirugía de hallazgos. "
            "NUNCA incluyas ANTECEDENTE DE, T4N3M1, EC IVA, ECOG, LIGHT, CAPRINI, PADUA, ROCKALL, "
            "ASA ni ninguna clasificación. Solo diagnósticos activos.\n"
            "3. Tercer párrafo: comienza con 'AL PASE DE VISITA SE ENCUENTRA EN LO NEUROLÓGICO ...' "
            "y ENUMERA en el mismo párrafo corrido los aparatos y sistemas EN ESTE ORDEN OBLIGATORIO: "
            "NEUROLÓGICO → RESPIRATORIO → CARDIOVASCULAR → GASTRO METABÓLICO → HÍDRICO URINARIO → "
            "HEMATOINFECCIOSO. Cada sistema con la fórmula 'EN LO [SISTEMA] ...' seguido de la "
            "información. No usar listas ni viñetas.\n"
            "4. Párrafo de exploración: comienza con 'A LA EXPLORACIÓN FÍSICA ...' en párrafo corrido, "
            "en el orden habitual (cabeza y cuello, tórax, abdomen, extremidades). Nunca escribas "
            "'SIN DATOS APORTADOS', 'SIN EVIDENCIA DOCUMENTADA', 'SIN HALLAZGOS REPORTADOS'.\n"
            "5. 'PLAN: ...' en UN SOLO renglón o máximo 2, separado por comas (ayuno/dieta, soluciones, "
            "analgesia IV, antibioticoterapia IV, rutina de UTI, vigilancia hemodinámica, "
            "cuantificación de drenajes/sondas, toma de laboratorios AM). NUNCA viñetas ni recomendaciones tipo ChatGPT.\n"
            "6. Cierra con 'PRONÓSTICO RESERVADO A EVOLUCIÓN.' en una línea sola.\n"
            "7. TODA la nota en MAYÚSCULAS. Sin listas, sin viñetas, sin numeración, sin markdown."
        )
    else:
        format_rules = (
            "FORMATO OBLIGATORIO (imitando los ejemplos permanentes de HOSPITALIZACIÓN del servicio):\n"
            "1. Inicia con 'NOTA DE EVOLUCIÓN POR CIRUGÍA GENERAL' en la línea 1 (o "
            "'NOTA DE EVOLUCIÓN POR CIRUGÍA GENERAL / [SUBESPECIALIDAD]' si aplica).\n"
            "2. Segundo párrafo — INTRODUCCIÓN EXACTA: 'SE TRATA DE PACIENTE MASCULINO/FEMENINO DE "
            "XX AÑOS, NOMBRE COMPLETO, CON DIAGNÓSTICO DE __________ POR LO QUE EL DÍA __________ "
            "SE LE REALIZÓ __________ CON HALLAZGOS DE __________.' NUNCA separes cirugía de hallazgos. "
            "NUNCA incluyas ANTECEDENTE DE, T4N3M1, EC IVA, ECOG, LIGHT, CAPRINI, PADUA, ROCKALL, "
            "ASA ni ninguna clasificación. Solo diagnósticos activos.\n"
            "3. Tercer párrafo: comienza con 'AL PASE DE VISITA SE ENCUENTRA ...' en un párrafo "
            "corrido que integra naturalmente: signos vitales INTERPRETADOS (nunca cifras crudas — "
            "usa 'SIGNOS VITALES DENTRO DE PARÁMETROS NORMALES', 'CON TENDENCIA A HIPERTENSIÓN', "
            "'CON TENDENCIA A HIPOTENSIÓN', 'AFEBRIL'), oxígeno si aplica ('CON PUNTAS NASALES A X LPM' "
            "o 'CON PUNTAS NASALES DE ALTO FLUJO A X LPM FIO2 X%' o 'CON VENTILACIÓN MECÁNICA' + "
            "'MANTENIENDO SATURACIÓN ADECUADA'), dolor SOLO con estas fórmulas: 'ADECUADO CONTROL "
            "ANALGÉSICO' / 'MAL CONTROL ANALGÉSICO, REFIERE ENA EN X, POR LO QUE SE INDICÓ X' / "
            "'PARCIAL CONTROL ANALGÉSICO, REFIERE ENA EN X, POR LO QUE SE INDICÓ X' (NUNCA 'DOLOR "
            "CONTROLADO'), dieta con abreviaturas del servicio SIN EXPANDIR (BOOST, NPT, NET), "
            "'SIN REFERIR NÁUSEAS O VÓMITO' (nunca 'SIN REFERENCIA DE...'), 'DIURESIS POR MICCIÓN "
            "ESPONTÁNEA...' o 'DIURESIS POR SONDA FOLEY DE CARACTERÍSTICAS CLARAS/HEMATÚRICAS "
            "CUANTIFICADA EN...' (integrar DKH si aplica), 'EVACUACIONES PRESENTES DE "
            "CARACTERÍSTICAS BRISTOL X' o 'EVACUACIONES PENDIENTES', drenajes con gasto y "
            "características, antibióticos si aplica.\n"
            "4. Laboratorios interpretados dentro del mismo párrafo o al final: NUNCA copies "
            "biometrías completas. Usa: 'DISMINUCIÓN DE LEUCOCITOS DE X A X', 'AUMENTO DE "
            "HEMOGLOBINA DE X A X', 'PERSISTE LEUCOCITOSIS (X MIL)', 'PERSISTE HIPOALBUMINEMIA (X)', "
            "'MEJORÍA DE FUNCIÓN RENAL', 'EMPEORAMIENTO DE FUNCIÓN RENAL'. Solo cambios relevantes.\n"
            "5. Cuarto párrafo: comienza con 'A LA EXPLORACIÓN FÍSICA ALERTA, REACTIVO, COOPERADOR' "
            "(o similar) y describe cabeza y cuello, cardiorrespiratorio, abdomen con herida y "
            "drenajes, extremidades. Un solo párrafo corrido. Nunca 'SIN DATOS APORTADOS', "
            "'SIN EVIDENCIA DOCUMENTADA', 'SIN HALLAZGOS REPORTADOS'.\n"
            "6. 'PLAN: ...' en UN SOLO renglón o máximo 2, separado por comas: "
            "'PLAN: AYUNO/DIETA X, ANALGESIA IV, ANTIBIÓTICOS IV, CUANTIFICACIÓN DE DRENAJES/SONDA, "
            "VIGILANCIA DE DATOS DE SANGRADO/DIFICULTAD RESPIRATORIA/ABDOMEN AGUDO'. NUNCA viñetas, "
            "NUNCA recomendaciones tipo ChatGPT.\n"
            "7. Cierra con 'PRONÓSTICO RESERVADO A EVOLUCIÓN.' en una línea sola.\n"
            "8. TODA la nota en MAYÚSCULAS. Sin listas, sin viñetas, sin numeración, sin markdown."
        )

    system = (
        "Eres RESIDENTE R3 DE CIRUGÍA GENERAL del HOSPITAL ÁNGELES redactando una nota MedSys "
        "para copiar y pegar. Debes escribir EXACTAMENTE con el estilo del servicio, no como un "
        "modelo de lenguaje.\n\n"
        "REGLAS DE ESTILO INAMOVIBLES:\n"
        "- TODO EL TEXTO EN MAYÚSCULAS (excepto números y símbolos).\n"
        "- SIN LISTAS. SIN VIÑETAS. SIN NUMERACIÓN. TODO EN PÁRRAFOS CORRIDOS separados por saltos de línea.\n"
        "- SIN MARKDOWN. SIN ASTERISCOS. SIN GUIONES DE LISTA. SIN 'ND'/'N/A' innecesarios.\n"
        "- NUNCA sonar como ChatGPT: sin frases de relleno, sin 'basado en la información', sin 'en resumen'.\n"
        "- NUNCA inventes datos, dosis, hallazgos, cultivos ni signos vitales que no estén presentes.\n"
        "- Español médico mexicano, léxico del servicio (BLAKE, TAMS, FIO2, VMI, RHP, ENA, EVN, etc.).\n"
        "- Longitud similar a los ejemplos: densa pero no exhaustiva.\n"
        "- Los ejemplos incluidos son SOLO REFERENCIA DE ESTILO: aprende su estructura, orden, tono, "
        "vocabulario y longitud. NUNCA copies nombres, diagnósticos, fechas ni frases textuales."
    )

    user_prompt = (
        f"{data_block}\n"
        f"TIPO DE NOTA A REDACTAR: {plantilla}\n\n"
        f"{format_rules}\n"
        f"{builtin_style_block}"
        f"{user_style_block}"
        f"\nAHORA REDACTA LA NOTA DEL PACIENTE INDICADO EN LA INFORMACIÓN FIJA. RECUERDA: "
        f"NUNCA copies texto de los ejemplos, solo IMITA el estilo. Devuelve SOLO la nota, sin comentarios."
    )

    note = await llm_generate(system, user_prompt)

    # WhatsApp — Prompt 5: solo primer nombre + (cama), sin edad/dx/DPQX/DEIH,
    # signos vitales interpretados, labs solo si cambian conducta, sin frases IA.
    wa_user_examples = _style_block(training.get("whatsapp_examples", ""), "MENSAJES DEL PROPIO USUARIO")
    wa_builtin_examples = whatsapp_style_block()

    # Extraer SOLO el primer nombre del paciente
    full_name = (patient.get("name") or "").strip()
    first_name = full_name.split()[0].title() if full_name else "________"
    patient_bed = (patient.get("bed") or "").strip()

    wa_prompt = (
        f"Redacta un mensaje de WhatsApp corto para el médico tratante como lo enviaría un R3 de "
        f"Cirugía General del Hospital Ángeles. Estilo natural, conversacional, breve. Un solo "
        f"párrafo corrido. Capitalización normal (NO mayúsculas).\n\n"
        f"DATOS INTERNOS (usa solo lo necesario, NO copies nombre completo, edad, diagnóstico, "
        f"DPQX ni DEIH en el mensaje):\n{ctx}\n\n"
        f"DATOS DEL DÍA ({target}):\n"
        f"- Dictado: {entry.get('dictation', '') or 'sin dictado.'}\n"
        f"- Labs: {entry.get('labs', '') or 'sin nuevos labs.'}\n"
        f"- Estudios: {entry.get('studies', '') or 'ninguno.'}\n"
        f"- Procedimientos: {entry.get('procedures', '') or 'ninguno.'}\n"
        f"- Cultivos: {entry.get('cultures', '') or 'ninguno.'}\n"
        f"- Signos vitales: {entry.get('vital_signs', '') or 'estables.'}\n"
        f"- Eventos: {entry.get('events', '') or 'sin eventos.'}\n\n"
        f"REGLAS OBLIGATORIAS:\n"
        f"- INICIA con esta fórmula exacta: 'Buenos días Dr, pasé a ver a {first_name} ({patient_bed}), está con...' "
        f"(o una variante muy cercana como 'Buenos días Dr, pase a ver a {first_name} ({patient_bed})...' / "
        f"'Hola doctor buenos días! Pasé a ver a {first_name} ({patient_bed})...').\n"
        f"- USA SOLO EL PRIMER NOMBRE ({first_name}) — NUNCA el nombre completo.\n"
        f"- NUNCA menciones edad, diagnóstico, DPQX ni DEIH.\n"
        f"- Interpreta signos vitales (NO copies cifras): 'signos vitales estables', "
        f"'con tendencia a hipertensión', 'con tendencia a hipotensión', 'afebril', "
        f"'con tendencia a taquicardia'. Solo agrega una cifra si hay un pico clínicamente importante.\n"
        f"- Dolor: 'refiere buen control analgésico' o 'refiere mal control analgésico...'. Nunca 'dolor controlado'.\n"
        f"- Dieta: 'tolera bien la dieta X' o 'en ayuno'. Abreviaturas BOOST/NPT/NET sin expandir.\n"
        f"- 'sin referir náuseas o vómito' (nunca 'sin referencia de').\n"
        f"- Diuresis: 'diuresis por micción espontánea' o 'diuresis por sonda Foley cuantificada en X'. Menciona drenajes con gasto si aplica ('Pleurex con gasto...').\n"
        f"- Laboratorios: solo menciónalos si cambian conducta. Ejemplos: 'leucocitos subieron a 41 mil', 'Hb bajó a 6.8', 'glucosa en 286 ya en manejo con glargina'. Nunca comentes labs normales.\n"
        f"- NUNCA uses: 'OJO', 'llama la atención', 'se sugiere', 'se recomienda', 'se informa que'.\n"
        f"- A veces (SOLO si hay decisión pendiente) termina con UNA sola pregunta: "
        f"'¿Gusta que retiremos el drenaje?', '¿Gusta que ajustemos la analgesia?', "
        f"'¿Gusta que iniciemos vía oral?', '¿Gusta que solicitemos TAC?', '¿Gusta que valoremos egreso?'. "
        f"Si NO hay decisión pendiente, cierra natural ('Quedo al pendiente, saludos!', 'Lindo día doc!', "
        f"'Saludos y bonito día').\n"
        f"- Un solo párrafo corrido. Sin listas. Sin viñetas. Sin markdown.\n"
        f"- Nunca inventes datos. Nunca suenes como IA.\n"
        f"{wa_builtin_examples}"
        f"{wa_user_examples}"
        f"\nDevuelve SOLO el mensaje de WhatsApp, sin comentarios."
    )
    system_wa = (
        "Eres residente R3 de Cirugía General del Hospital Ángeles enviando WhatsApp corto al "
        "médico tratante. Solo primer nombre y cama del paciente — NUNCA nombre completo, edad, "
        "diagnóstico, DPQX ni DEIH. Interpretas signos vitales y labs (no cifras crudas salvo "
        "picos relevantes). Un solo párrafo corrido, capitalización normal. Nunca sonar a IA. "
        "Nunca decir 'OJO', 'llama la atención', 'se sugiere', 'se recomienda'."
    )
    wa = await llm_generate(system_wa, wa_prompt)

    await db.daily_entries.update_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": target},
        {"$set": {"ai_evolution_note": note, "ai_whatsapp": wa, "updated_at": now_utc().isoformat()}},
        upsert=True,
    )
    return {"note": note, "whatsapp": wa}


@api.post("/pase/today")
async def generate_full_pase(user: dict = Depends(get_user)):
    target = today_iso()
    patients = await db.patients.find({"user_id": user["id"], "active": True}, {"_id": 0}).to_list(500)
    patients.sort(key=lambda x: ((x.get("floor") or ""), (x.get("bed") or "")))

    sections = []
    for p in patients:
        entry = await db.daily_entries.find_one(
            {"patient_id": p["id"], "user_id": user["id"], "date": target}, {"_id": 0}
        )
        summary = entry.get("ai_pase_summary") if entry else None
        if summary:
            sections.append(summary.strip())

    header = f"PASE DE VISITA — {target}\nCirugía General\n{'=' * 40}\n"
    body = "\n\n" + ("\n\n" + ("-" * 40) + "\n\n").join(sections) if sections else "\n\n(No hay pacientes con pase generado hoy.)\n"
    return {"pase": header + body, "patient_count": len(sections)}


# ---------------- Discharged patients ----------------
@api.post("/patients/{patient_id}/readmit")
async def readmit_patient(patient_id: str, user: dict = Depends(get_user)):
    result = await db.patients.update_one(
        {"id": patient_id, "user_id": user["id"]},
        {"$set": {"active": True, "discharged_at": None, "updated_at": now_utc().isoformat()}},
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    return {"ok": True}


# ---------------- CENSO IMPORT (Word .docx) ----------------
CENSO_FIELDS = [
    "name", "age", "sex", "bed", "floor", "service",
    "attending_physician", "admission_date",
    "surgery_date", "dx_short", "dx_full", "surgery_procedure",
    "surgery_findings", "medical_history", "allergies",
    "important_medications", "oncology_treatment",
    "oncology_status", "unit_classification",
    "is_surgical", "is_pending_discharge",
]

CENSO_DAILY_FIELDS = ["labs", "studies", "procedures", "cultures", "events"]


def _extract_docx_text(content: bytes) -> str:
    """Extract text from .docx preserving TABLE STRUCTURE:
    each row is emitted as one block, each cell labeled [COL N] with the full
    multi-paragraph content of the cell. This lets GPT identify the 5th column
    reliably so it can classify labs / imaging / procedures / cultures.
    """
    buf = io.BytesIO(content)
    doc = DocxDocument(buf)
    parts: list[str] = []
    # Paragraphs (headers, notes outside tables)
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Tables: preserve column structure
    for t_idx, table in enumerate(doc.tables):
        parts.append(f"\n===== TABLA {t_idx + 1} =====")
        for r_idx, row in enumerate(table.rows):
            row_has_content = any(c.text.strip() for c in row.cells)
            if not row_has_content:
                continue
            parts.append(f"\n----- FILA {r_idx + 1} -----")
            for c_idx, cell in enumerate(row.cells):
                # Preserve internal paragraphs of the cell (line breaks matter for column 5)
                lines = [p.text for p in cell.paragraphs if p.text.strip()]
                if not lines:
                    continue
                cell_text = "\n".join(lines)
                parts.append(f"[COL {c_idx + 1}]:\n{cell_text}")
    return "\n".join(parts)


# ---------------- Column-5 regex classifier (fallback + reinforcement) ----------------
# STRICT keyword lists per user spec (Etapa 3):
# - Labs: SOLO LABS, GASA, GASV, EGO (con su contenido). No absorber TAC/RHP/Cultivos.
_LABS_KW = _re.compile(
    r"(?:^|\W)(LABS|GASA|GASV|EGO)\b",
    _re.IGNORECASE,
)
# - Estudios de imagen: RX, RXTX, TAC (con variantes), ANGIOTAC, RM/RMN, USG,
#   PET-CT, ECG, ECOCARDIOGRAMA, SEGD, MASTOGRAFIA. NUNCA RHP/Colonoscopia/Cirugía.
_IMAGING_KW = _re.compile(
    r"(?:^|\W)(RXTX|RX|TAC(?:\s+(?:ABD|TX|C\/C|IV|VO|SIMPLE|CONTRAST))?|ANGIOTAC|"
    r"RM|RMN|USG|ULTRASONIDO|PET[\-\s]?CT|ECG|EKG|ECOCARDIOGRAMA|ECO|"
    r"SEGD|SERIE\s+ESOFAGOGA|MASTOGRAF|RADIOGRAF|TOMOGRAF|RESONANCIA)\b",
    _re.IGNORECASE,
)
# - Procedimientos: Colonoscopia, Endoscopia, Panendoscopia, Cirugía, Hallazgos,
#   Biopsias, RHP, Paracentesis, Drenajes, Toracocentesis. Y procedimientos QX con fecha.
_PROC_KW = _re.compile(
    r"(?:^|\W)(PANENDOSCOPIA|COLONOSCOPIA|ENDOSCOPIA|RHP|HALLAZGOS(?:\s+QUIR)?|"
    r"DRENAJES?|PARACENTESIS|TORACOCENTESIS|BIOPSIAS?|CIRUG(?:I|Í)A|QUIR[UÚ]RGIC|"
    r"HEMICOLECTOM(?:I|Í)A|COLECISTECTOM(?:I|Í)A|APENDICECTOM(?:I|Í)A|"
    r"LAPAROSCOP|LAPE|LAPAROTOM|ITALLMR|ILEOSTOM|COLOSTOM|GASTROSTOM|"
    r"YEYUNOSTOM|ANASTOMOSIS|RESECC[IÓ]N|ANEXECTOM|HISTERECTOM|OOFERECTOM|"
    r"MASTECTOM|ADENECTOM|SPLENECTOM|POST[\s\-]?OPERATORIO|POP\b|PO\s+LAPE)\b",
    _re.IGNORECASE,
)
# - Cultivos: Cultivos, Hemocultivos, Urocultivos, Gram, Susceptibilidad, BLEE,
#   Microorganismos, Antibiograma, Desarrollo, Aislamiento.
_CULT_KW = _re.compile(
    r"(?:^|\W)(HEMOCULTIVOS?|UROCULTIVOS?|CULTIVOS?|GRAM|SUSCEPTIBILIDAD|"
    r"DESARROLLO|BLEE|MICROORGANISMOS?|AISLAMIENTO|ANTIBIOGRAMA)\b",
    _re.IGNORECASE,
)

# Clasificaciones / escalas (van al final del dx_full como CLASIFICACIONES:)
# NO son diagnósticos ni procedimientos.
_CLASSIFICATION_KW = _re.compile(
    r"\b("
    r"PADUA\s*\d+|"
    r"ROCKALL\s*\d+|"
    r"ECOG\s*[0-4]|"
    r"KARNOFSKY\s*\d+|"
    r"CHILD[\s\-]?PUGH\s*[A-C]|"
    r"MELD\s*\d+|"
    r"APACHE\s*(?:II|III|IV)?\s*\d*|"
    r"SOFA\s*\d+|"
    r"GLASGOW\s*\d+|GCS\s*\d+|"
    r"NIHSS\s*\d+|"
    r"EC\s*[IV]+[A-C]?|"                    # Estadio clínico romano
    r"ESTADIO\s*[IV]+[A-C]?|"
    r"T[0-4]N[0-3]M[0-1X]|"                  # TNM
    r"TNM\s*T[0-4]N[0-3]M[0-1X]|"
    r"BCLC\s*[0-D]|"
    r"OKUDA\s*[I-V]+|"
    r"BREUSH|"
    r"BAKER\s*[I-V]+|"
    r"AAST\s*[I-V]+|"
    r"HINCHEY\s*[I-V]+[a-c]?|"
    r"MRC\s*[0-5]|"
    r"MRC\s*DYSPNEA|"
    r"WELLS\s*[\d\.]+|"
    r"GENEVA\s*[\d\.]+|"
    r"CHA2DS2[\s\-]?VASC\s*\d+|"
    r"HAS[\s\-]?BLED\s*\d+"
    r")\b",
    _re.IGNORECASE,
)


def _classify_col5_text(text: str) -> dict:
    """Split col 5 text into labs/studies/procedures/cultures.
    STRICT rules per user spec (Etapa 3):
    - labs: SOLO chunks que empiezan con LABS/GASA/GASV/EGO
    - studies: SOLO chunks que empiezan con RX/RXTX/TAC/ANGIOTAC/RM/RMN/USG/PET/ECG/ECO/SEGD/MASTOGRAF/etc.
    - procedures: SOLO chunks que empiezan con COLONOSCOPIA/ENDOSCOPIA/PANENDOSCOPIA/CIRUGIA/RHP/HALLAZGOS/BIOPSIAS/PARACENTESIS/DRENAJES/etc.
    - cultures: SOLO chunks que empiezan con CULTIVOS/HEMOCULTIVOS/UROCULTIVOS/GRAM/BLEE/etc.
    - Un chunk termina cuando comienza otro keyword de OTRA categoría.
    """
    if not text or not text.strip():
        return {"labs": "", "studies": "", "procedures": "", "cultures": "", "events": ""}
    raw = text.strip()

    # Anchor: chunk boundary starts BEFORE any category keyword OR a date.
    # Ordering matters: more specific first so ANGIOTAC beats TAC etc.
    _SPLIT_ANCHOR = _re.compile(
        r"(?=(?:^|\n|\.\s+|;\s*)"
        r"(?:LABS|GASA|GASV|EGO|"
        r"ANGIOTAC|TAC(?:\s+(?:ABD|TX|C\/C|IV|VO|SIMPLE|CONTRAST))?|"
        r"RXTX|RX|RMN|RM|USG|ULTRASONIDO|"
        r"PET[\-\s]?CT|ECG|EKG|ECOCARDIOGRAMA|ECO|"
        r"SEGD|MASTOGRAF|RADIOGRAF|TOMOGRAF|RESONANCIA|"
        r"PANENDOSCOPIA|COLONOSCOPIA|ENDOSCOPIA|RHP|HALLAZGOS|"
        r"CIRUG(?:I|Í)A|QUIR[UÚ]RGIC|PO\s+LAPE|LAPE|LAPAROSCOP|LAPAROTOM|"
        r"HEMICOLECTOM|COLECISTECTOM|APENDICECTOM|ITALLMR|ILEOSTOM|COLOSTOM|"
        r"PARACENTESIS|TORACOCENTESIS|BIOPSIAS?|DRENAJES?|"
        r"HEMOCULTIVOS?|UROCULTIVOS?|CULTIVOS?|GRAM|BLEE|"
        r"SUSCEPTIBILIDAD|DESARROLLO|MICROORGANISMOS?|ANTIBIOGRAMA|AISLAMIENTO)"
        r"\b)",
        _re.IGNORECASE,
    )
    # Also split before standalone dates (a "fecha subrayada" starts a new item)
    _DATE_SPLIT = _re.compile(r"(?=\n\s*\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b)")

    # First split by keyword anchors
    chunks = _SPLIT_ANCHOR.split(raw)
    # Then refine by dates
    refined: list[str] = []
    for ch in chunks:
        ch = ch.strip()
        if not ch:
            continue
        subs = _DATE_SPLIT.split(ch)
        for s in subs:
            s = s.strip()
            if s:
                refined.append(s)

    if not refined:
        refined = [raw]

    bins = {"labs": [], "studies": [], "procedures": [], "cultures": [], "events": []}

    def _classify_chunk(chunk: str) -> str:
        """Return the category based on the FIRST keyword found (leading token wins)."""
        # Strip any leading date to look at the actual keyword
        head = _re.sub(r"^\s*\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\s*[\n\-]?\s*", "", chunk, count=1)
        head_upper = head.upper().lstrip()

        # Check cultures first (very specific)
        if _re.match(r"(?:HEMOCULTIVOS?|UROCULTIVOS?|CULTIVOS?|GRAM|BLEE|SUSCEPTIBILIDAD|DESARROLLO|MICROORGANISMOS?|ANTIBIOGRAMA|AISLAMIENTO)\b", head_upper):
            return "cultures"
        # Procedures next
        if _re.match(r"(?:PANENDOSCOPIA|COLONOSCOPIA|ENDOSCOPIA|RHP|HALLAZGOS|CIRUG|QUIR|PO\s+LAPE|LAPE|LAPAROSCOP|LAPAROTOM|HEMICOLECTOM|COLECISTECTOM|APENDICECTOM|ITALLMR|ILEOSTOM|COLOSTOM|PARACENTESIS|TORACOCENTESIS|BIOPSIAS?|DRENAJES?)\b", head_upper):
            return "procedures"
        # Imaging studies
        if _re.match(r"(?:ANGIOTAC|TAC|RXTX|RX|RMN|RM|USG|ULTRASONIDO|PET|ECG|EKG|ECOCARDIOGRAMA|ECO|SEGD|MASTOGRAF|RADIOGRAF|TOMOGRAF|RESONANCIA)\b", head_upper):
            return "studies"
        # Labs
        if _re.match(r"(?:LABS|GASA|GASV|EGO)\b", head_upper):
            return "labs"
        return "events"

    for chunk in refined:
        cat = _classify_chunk(chunk)
        bins[cat].append(chunk)

    return {k: "\n\n".join(v) for k, v in bins.items()}


async def _parse_censo_with_gpt(raw_text: str) -> list[dict]:
    """DEPRECATED: kept for compatibility, no longer called. Census parsing is now 100% deterministic."""
    raise RuntimeError("_parse_censo_with_gpt is deprecated. Use _parse_censo_docx_deterministic.")


# ---------------- 100% deterministic .docx census parser (NO GPT) ----------------
_NAME_RE = _re.compile(r"[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ'\-]{1,}(?:\s+[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ'\-]{1,}){1,}")
_AGE_RE = _re.compile(r"(\d{1,3})\s*(?:años?|a\b|A\b|Y|y/o|yo)", _re.IGNORECASE)
_SEX_RE = _re.compile(r"(?:^|\W)(masculino|femenino|hombre|mujer|\bM\b|\bF\b)", _re.IGNORECASE)
_BED_RE = _re.compile(r"\b(S\d{2,3}|\d{3})\b")
_ATTENDING_RE = _re.compile(
    r"(?:Tratante|Adscrito|M[ée]dico responsable)\s*[:\-]?\s*(Dr(?:a)?\.?\s+[A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑñáéíóúü'\-\.\s]{1,60})",
    _re.IGNORECASE,
)
_ATTENDING_LOOSE_RE = _re.compile(r"\b(Dr(?:a)?\.?\s+[A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑñáéíóúü'\-\.]{1,40}(?:\s+[A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑñáéíóúü'\-\.]{1,40})?)\b")
_DATE_RE = _re.compile(r"\b(\d{1,2})[/\-](\d{1,2})[/\-](\d{2,4})\b")
_ALLERGY_RE = _re.compile(r"(?:Alergias?|APNP)\s*[:\-]\s*([^\n]{1,120})", _re.IGNORECASE)
_MEDS_RE = _re.compile(r"(?:Medicamentos?|Meds?)\s*[:\-]\s*([^\n]{1,200})", _re.IGNORECASE)
_APP_RE = _re.compile(r"(?:APP|Antecedentes(?:\s+personales)?(?:\s+patol[óo]gicos)?)\s*[:\-]\s*([^\n]{1,300})", _re.IGNORECASE)
_ONCO_RE = _re.compile(r"(?:Oncol[óo]gico|Estadio|Etapa)\s*[:\-]?\s*([^\n]{1,120})", _re.IGNORECASE)
_UTI_RE = _re.compile(r"\b(UTI|UTIM|Terapia\s+Intensiva|Terapia\s+Intermedia)\b", _re.IGNORECASE)
_PENDING_DISCHARGE_RE = _re.compile(r"(alta\s+pendiente|alta\s+hoy|candidato\s+a\s+egreso|en\s+espera\s+de\s+alta)", _re.IGNORECASE)


def _iso_date(s: str) -> Optional[str]:
    m = _DATE_RE.search(s)
    if not m:
        return None
    d, mo, y = m.group(1), m.group(2), m.group(3)
    if len(y) == 2:
        y = "20" + y
    try:
        return f"{int(y):04d}-{int(mo):02d}-{int(d):02d}"
    except Exception:
        return None


def _cell_classification_score(text: str) -> int:
    return (
        len(_LABS_KW.findall(text)) + len(_IMAGING_KW.findall(text))
        + len(_PROC_KW.findall(text)) + len(_CULT_KW.findall(text))
    )


def _parse_row_deterministic(cells: list[str]) -> Optional[dict]:
    """POSITIONAL parser (fixed layout, 1-indexed like user spec):
    COL 1 (idx 0): bed + tratante + residentes
    COL 2 (idx 1): name + edad + DEIH + DPQX
    COL 3 (idx 2): DIAGNÓSTICOS (ONLY diagnoses, never indications)
    COL 4 (idx 3): INDICACIONES médicas (goes to important_medications, NEVER to dx)
    COL 5 (idx 4): labs / gabinete (classified into labs/studies/procedures/cultures)
    COL 6 (idx 5): signos vitales
    """
    if not any(c.strip() for c in cells):
        return None
    while len(cells) < 6:
        cells.append("")

    p: dict = {}

    # ---- COL 1 (idx 0): bed + attending + residents ----
    col0_lines = [ln.strip() for ln in cells[0].split("\n") if ln.strip()]
    if col0_lines:
        m = _BED_RE.search(col0_lines[0])
        if m:
            p["bed"] = m.group(1).upper()
        if len(col0_lines) >= 2:
            p["attending_physician"] = col0_lines[1].strip()

    # ---- COL 2 (idx 1): name + age (validation) ----
    col1_lines = [ln.strip() for ln in cells[1].split("\n") if ln.strip()]
    if not col1_lines:
        return None
    raw_name = col1_lines[0].strip()
    if _re.search(r"\bDRA?\.?\s*$", raw_name, _re.IGNORECASE):
        raise ValueError(
            f"Alineamiento de columnas perdido: nombre '{raw_name}' termina en DR/DRA. Se cancela la importación."
        )
    p["name"] = raw_name
    if len(col1_lines) >= 2:
        m = _re.search(r"\b(\d{1,3})\b", col1_lines[1])
        if m:
            try: p["age"] = int(m.group(1))
            except Exception: pass

    # ---- COL 3 (idx 2): DIAGNÓSTICOS + separación de PROCEDIMIENTOS y CLASIFICACIONES ----
    # Reglas Etapa 3:
    #  - Diagnósticos reales quedan en dx_short/dx_full.
    #  - Procedimientos con fecha (12/06/26 PANENDOSCOPIA, PO LAPE + HEMICOLECTOMIA, etc.)
    #    se extraen y se mueven a daily_entry.procedures.
    #  - Clasificaciones (PADUA 7, ECOG 1, T4N3M1, ROCKALL 6, EC IVA...) se agregan al
    #    final del dx_full bajo el subtítulo "CLASIFICACIONES:", pero NO en dx_short.
    _DRUG_RE = _re.compile(
        r"(?:\d+\s*(?:mg|g|mcg|mcgs|ml|ml/h|ml/hr|UI|meq|gtas?|comp|caps|cc)\b"
        r"|c/\d+\s*h\b"
        r"|\b(?:IV|VO|SL|IM|SC|EV|PO)\b"
        r"|\b(?:PRN|c/8|c/12|c/24|c/6)\b)",
        _re.IGNORECASE,
    )
    _DATE_LINE_RE = _re.compile(r"^\s*\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\s*$")
    _PROC_IN_DX_RE = _re.compile(
        r"^\s*(?:\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\s+)?"
        r"(?:PANENDOSCOPIA|COLONOSCOPIA|ENDOSCOPIA|RHP|PO\s+LAPE|LAPE|"
        r"HEMICOLECTOM|COLECISTECTOM|APENDICECTOM|ITALLMR|LAPAROSCOP|LAPAROTOM|"
        r"HALLAZGOS|PARACENTESIS|TORACOCENTESIS|BIOPSIAS?|DRENAJES?|"
        r"ILEOSTOM|COLOSTOM|GASTROSTOM|YEYUNOSTOM|ANASTOMOSIS|RESECC[IÓ]N|"
        r"ANEXECTOM|HISTERECTOM|OOFERECTOM|MASTECTOM|SPLENECTOM)",
        _re.IGNORECASE,
    )

    dx_cell = cells[2].strip()
    dx_lines_raw = [ln.strip() for ln in dx_cell.split("\n") if ln.strip()]

    dx_clean: list[str] = []
    proc_lines: list[str] = []
    class_lines: list[str] = []
    pending_date: Optional[str] = None  # a date line waits for its next content line

    for line in dx_lines_raw:
        # Pure date line — retain for the next content line
        if _DATE_LINE_RE.match(line):
            pending_date = line
            continue

        # Skip pure DEIH/DPQX/number-only
        if _re.fullmatch(r"\s*(DEIH|DPQX)?\s*\d*\s*", line, _re.IGNORECASE):
            pending_date = None
            continue

        # Classification-only line: capture and skip
        if _CLASSIFICATION_KW.search(line) and not _PROC_IN_DX_RE.match(line):
            # If the line is ONLY classifications, redirect entirely
            # e.g. "PADUA 7", "ECOG 1", "T4N3M1", "EC IVA"
            stripped = line.strip()
            # If the line has some other diagnosis words + a classification,
            # we still keep the whole line as dx (it's a diagnosis WITH a classification).
            # Heuristic: if line length is short (<= 25 chars) and matches only classification pattern, treat as pure classification.
            only_class = _re.fullmatch(
                r"\s*(?:PADUA\s*\d+|ROCKALL\s*\d+|ECOG\s*[0-4]|KARNOFSKY\s*\d+|"
                r"CHILD[\s\-]?PUGH\s*[A-C]|MELD\s*\d+|APACHE\s*(?:II|III|IV)?\s*\d*|"
                r"SOFA\s*\d+|GLASGOW\s*\d+|GCS\s*\d+|NIHSS\s*\d+|"
                r"EC\s*[IV]+[A-C]?|ESTADIO\s*[IV]+[A-C]?|"
                r"T[0-4]N[0-3]M[0-1X]|TNM\s*T[0-4]N[0-3]M[0-1X]|"
                r"BCLC\s*[0-D]|OKUDA\s*[I-V]+|BAKER\s*[I-V]+|AAST\s*[I-V]+|"
                r"HINCHEY\s*[I-V]+[a-c]?|MRC\s*[0-5]|WELLS\s*[\d\.]+|"
                r"GENEVA\s*[\d\.]+|CHA2DS2[\s\-]?VASC\s*\d+|HAS[\s\-]?BLED\s*\d+)\s*",
                stripped,
                _re.IGNORECASE,
            )
            if only_class:
                class_lines.append(stripped.upper())
                pending_date = None
                continue
            # otherwise fall-through to normal dx handling but capture also the classification
            found = _CLASSIFICATION_KW.findall(stripped)
            for f in found:
                if f.upper() not in class_lines:
                    class_lines.append(f.upper())
            # Continue processing the line (it may still contain a real dx)

        # Procedure line (with or without date prefix) → daily_entry.procedures
        if _PROC_IN_DX_RE.match(line):
            proc_text = line
            if pending_date:
                proc_text = f"{pending_date} {line}"
                pending_date = None
            proc_lines.append(proc_text)
            continue

        # Imaging / labs / cultures accidentally in dx column → skip (belongs to col5)
        if _cell_classification_score(line) > 0:
            pending_date = None
            continue

        # Drug/indication that leaked in → skip
        if _DRUG_RE.search(line):
            pending_date = None
            continue

        # Real diagnosis line
        dx_line = line
        if pending_date:
            # A date preceding a diagnosis is unusual, but keep it attached
            dx_line = f"{pending_date} {line}"
            pending_date = None
        dx_clean.append(dx_line)

    if dx_clean:
        p["dx_short"] = dx_clean[0][:200]
        dx_full_text = "\n".join(dx_clean)
        if class_lines:
            dx_full_text += "\n\nCLASIFICACIONES:\n" + "\n".join(class_lines)
        p["dx_full"] = dx_full_text[:2500]
    elif class_lines:
        # Only classifications, no explicit dx → store classifications in dx_full for reference
        p["dx_full"] = "CLASIFICACIONES:\n" + "\n".join(class_lines)

    if proc_lines:
        # Stash for daily_entry.procedures downstream merge
        p["_col3_procedures"] = "\n".join(proc_lines)

    # ---- COL 4 (idx 3): INDICACIONES → important_medications (NEVER dx) ----
    ind_text = cells[3].strip()
    if ind_text:
        # Store raw as important_medications (user request) AND stash for diff
        p["important_medications"] = ind_text[:4000]
        p["_indications_text"] = ind_text  # ephemeral: used for indications_changes diff

    # ---- COL 5 (idx 4): labs/studies/procedures/cultures ----
    col4 = cells[4].strip()
    if col4:
        buckets = _classify_col5_text(col4)
        p["labs"] = buckets["labs"]
        p["studies"] = buckets["studies"]
        # Merge procedures from Col3 (dates+QX lines) with Col5 procedures
        col5_procs = buckets["procedures"]
        col3_procs = p.pop("_col3_procedures", "")
        merged_procs = "\n\n".join(x for x in (col3_procs, col5_procs) if x)
        p["procedures"] = merged_procs
        p["cultures"] = buckets["cultures"]
        if buckets["events"]:
            p["events"] = buckets["events"]
    else:
        # Even if Col5 empty, still surface Col3-extracted procedures
        col3_procs = p.pop("_col3_procedures", "")
        if col3_procs:
            p["procedures"] = col3_procs

    # ---- COL 6 (idx 5): signos vitales ----
    col5 = cells[5].strip()
    if col5:
        p["vital_signs"] = col5

    # Metadata defaults
    p["unit_classification"] = "Piso"
    p["is_surgical"] = True
    p["is_pending_discharge"] = False

    return p

def _parse_censo_docx_deterministic(content: bytes) -> list[dict]:
    """Parse .docx census PURELY with python-docx + regex. NO GPT calls."""
    buf = io.BytesIO(content)
    doc = DocxDocument(buf)
    patients: list[dict] = []
    for table in doc.tables:
        rows = list(table.rows)
        if not rows:
            continue
        # Skip header row if it contains labels like "Nombre" / "Cama" / "Diagnóstico"
        start = 0
        header = " ".join(c.text for c in rows[0].cells).lower()
        if any(kw in header for kw in ("nombre", "paciente", "cama", "diagn", "cirug", "tratante", "servicio")):
            start = 1
        for row in rows[start:]:
            cells = []
            for cell in row.cells:
                # Preserve per-paragraph newlines within a cell
                lines = [p.text for p in cell.paragraphs if p.text.strip()]
                cells.append("\n".join(lines))
            try:
                p = _parse_row_deterministic(cells)
            except Exception:
                logger.exception("Deterministic parse failed on row")
                continue
            if p and p.get("name"):
                # DEBUG: print raw columns before saving (per user request)
                logger.info(
                    "CENSO ROW → name=%r bed=%r\n[COL1]=%r\n[COL2]=%r\n[COL4]=%r\n[COL5]=%r\n[COL6]=%r",
                    p.get("name"), p.get("bed"),
                    cells[0] if len(cells) > 0 else "",
                    cells[1] if len(cells) > 1 else "",
                    cells[3] if len(cells) > 3 else "",
                    cells[4] if len(cells) > 4 else "",
                    cells[5] if len(cells) > 5 else "",
                )
                logger.info(
                    "  PARSED → dx_short=%r  labs=%d studies=%d procedures=%d cultures=%d vital=%d",
                    p.get("dx_short"),
                    len(p.get("labs", "") or ""),
                    len(p.get("studies", "") or ""),
                    len(p.get("procedures", "") or ""),
                    len(p.get("cultures", "") or ""),
                    len(p.get("vital_signs", "") or ""),
                )
                patients.append(p)
    return patients


def _norm_name(s: str) -> str:
    import unicodedata
    if not s:
        return ""
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
    return " ".join(s.lower().strip().split())


# In-memory job store for async .docx imports (per-process, sufficient for single-user app)
IMPORT_JOBS: dict[str, dict] = {}


@api.post("/patients/import-censo")
async def import_censo(
    file: UploadFile = File(...),
    confirm: bool = False,
    user: dict = Depends(get_user),
):
    """Kick off async census import job. Parsing is 100% DETERMINISTIC (no GPT)."""
    filename = (file.filename or "").lower()
    if not filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="El archivo debe ser un .docx")
    try:
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Archivo vacío")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo leer el .docx: {e}")

    job_id = str(uuid.uuid4())
    IMPORT_JOBS[job_id] = {
        "status": "running",
        "user_id": user["id"],
        "started_at": now_utc().isoformat(),
        "confirm": confirm,
        "result": None,
        "error": None,
    }
    asyncio.create_task(_run_import_job(job_id, content, user["id"], confirm))
    return {"job_id": job_id, "status": "running", "confirm": confirm}


@api.get("/patients/import-censo/status/{job_id}")
async def import_censo_status(job_id: str, user: dict = Depends(get_user)):
    job = IMPORT_JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Trabajo no encontrado")
    if job["user_id"] != user["id"]:
        raise HTTPException(status_code=403, detail="No autorizado")
    return {
        "status": job["status"],
        "result": job.get("result"),
        "error": job.get("error"),
        "confirm": job.get("confirm"),
        "started_at": job.get("started_at"),
    }


async def _run_import_job(job_id: str, content_bytes: bytes, user_id: str, confirm: bool):
    """Background task that parses the census with the DETERMINISTIC parser (no GPT) and (optionally) upserts."""
    try:
        errors: list[str] = []
        try:
            parsed = _parse_censo_docx_deterministic(content_bytes)
        except ValueError as e:
            # Alignment error → abort entire import per user spec
            logger.error("Alignment error for job %s: %s", job_id, e)
            IMPORT_JOBS[job_id] = {
                **IMPORT_JOBS[job_id],
                "status": "failed",
                "error": f"{str(e)} No se creó ni actualizó ningún paciente.",
            }
            return
        except Exception as e:
            logger.exception("Deterministic parsing failed for job %s", job_id)
            IMPORT_JOBS[job_id] = {
                **IMPORT_JOBS[job_id],
                "status": "failed",
                "error": f"Error al leer el .docx: {str(e)[:300]}",
            }
            return
        if not parsed:
            IMPORT_JOBS[job_id] = {
                **IMPORT_JOBS[job_id],
                "status": "failed",
                "error": "No se detectaron tablas de pacientes en el .docx.",
            }
            return

        # Safety: even if _parse_row_deterministic missed it, double-check no name is a tratante
        for row in parsed:
            n = row.get("name", "")
            if _re.search(r"\bDRA?\.?\s*$", n, _re.IGNORECASE):
                IMPORT_JOBS[job_id] = {
                    **IMPORT_JOBS[job_id],
                    "status": "failed",
                    "error": f"Alineamiento de columnas perdido en la fila '{n}'. Importación cancelada.",
                }
                return

        existing = await db.patients.find({"user_id": user_id, "active": True}, {"_id": 0}).to_list(2000)
        name_to_patient = {_norm_name(p["name"]): p for p in existing}
        seen_names: set[str] = set()

        new_ones: list[dict] = []
        updated_ones: list[dict] = []
        today = today_iso()

        for order_idx, entry in enumerate(parsed):
            try:
                if not isinstance(entry, dict):
                    errors.append(f"Registro inválido (no es objeto): {str(entry)[:80]}")
                    continue
                name = (entry.get("name") or "").strip()
                if not name:
                    errors.append("Registro sin nombre — omitido.")
                    continue

                # ---- Column-5 safety net: guarantee classification ----
                # 1) If GPT accidentally dumped clinical bulk into dx_short/dx_full/events,
                #    or if any of labs/studies/procedures/cultures came empty, run
                #    the regex classifier on whatever loose text is available.
                bulk_sources = []
                for src_key in ("events", "dx_full", "surgery_findings"):
                    val = entry.get(src_key)
                    if isinstance(val, str) and val and len(val) > 40:
                        # Only reclassify if it looks like it has categorical keywords
                        if (_LABS_KW.search(val) or _IMAGING_KW.search(val)
                                or _PROC_KW.search(val) or _CULT_KW.search(val)):
                            bulk_sources.append((src_key, val))

                # Reclassify bulk_sources content into the 4 daily buckets
                for src_key, val in bulk_sources:
                    classified = _classify_col5_text(val)
                    for bucket in ("labs", "studies", "procedures", "cultures"):
                        added = classified.get(bucket, "")
                        if not added:
                            continue
                        existing = (entry.get(bucket) or "").strip()
                        if added.strip() not in existing:
                            entry[bucket] = (existing + "\n" + added).strip() if existing else added
                    # Remove the bulk from the source: only keep true "events" (non-classifiable)
                    entry[src_key] = classified.get("events", "").strip() if src_key == "events" else ""
                    # For dx_full/surgery_findings, keep only what regex could NOT classify (real dx text)
                    if src_key in ("dx_full", "surgery_findings"):
                        entry[src_key] = classified.get("events", "").strip()

                # 2) Even if GPT filled labs/studies/etc., re-run classifier ONLY on `events`
                #    (in case GPT dumped stuff there). This is idempotent.
                if entry.get("events"):
                    ev = entry["events"]
                    if (_LABS_KW.search(ev) or _IMAGING_KW.search(ev)
                            or _PROC_KW.search(ev) or _CULT_KW.search(ev)):
                        classified = _classify_col5_text(ev)
                        for bucket in ("labs", "studies", "procedures", "cultures"):
                            added = classified.get(bucket, "")
                            if not added:
                                continue
                            existing = (entry.get(bucket) or "").strip()
                            if added.strip() not in existing:
                                entry[bucket] = (existing + "\n" + added).strip() if existing else added
                        entry["events"] = classified.get("events", "").strip()

                key = _norm_name(name)
                seen_names.add(key)
                payload = {k: entry.get(k) for k in CENSO_FIELDS if entry.get(k) not in (None, "")}
                if "age" in payload:
                    try:
                        payload["age"] = int(payload["age"])
                    except Exception:
                        payload.pop("age", None)
                # Always update the censo_order for each imported patient
                payload["censo_order"] = order_idx

                prev = name_to_patient.get(key)
                if prev:
                    update = {k: v for k, v in payload.items() if v not in (None, "")}
                    if update:
                        update["updated_at"] = now_utc().isoformat()
                        if confirm:
                            await db.patients.update_one(
                                {"id": prev["id"], "user_id": user_id}, {"$set": update}
                            )
                        merged = {**prev, **update}
                        updated_ones.append({"id": prev["id"], "name": merged["name"]})
                    # Store labs/studies/events of the day (5th column) if present
                    if confirm:
                        await _upsert_daily_from_censo(prev["id"], user_id, today, entry)
                else:
                    if confirm:
                        new_patient = Patient(
                            user_id=user_id,
                            name=name,
                            is_new_admission=True,
                            **{k: v for k, v in payload.items() if k != "name"},
                        )
                        await db.patients.insert_one(new_patient.model_dump())
                        await _upsert_daily_from_censo(new_patient.id, user_id, today, entry)
                        new_ones.append({"id": new_patient.id, "name": new_patient.name})
                    else:
                        new_ones.append({"id": None, "name": name})
            except Exception as e:
                logger.exception("Failed processing entry %s", entry)
                errors.append(f"Error procesando '{entry.get('name', '?')}': {str(e)[:120]}")

        discharged_ones: list[dict] = []
        for key, p in name_to_patient.items():
            if key not in seen_names:
                try:
                    if confirm:
                        await db.patients.update_one(
                            {"id": p["id"], "user_id": user_id},
                            {"$set": {"active": False, "discharged_at": today_iso(), "updated_at": now_utc().isoformat()}},
                        )
                    discharged_ones.append({"id": p["id"], "name": p["name"]})
                except Exception as e:
                    errors.append(f"Error al mover a egresados '{p.get('name')}': {str(e)[:120]}")

        IMPORT_JOBS[job_id] = {
            **IMPORT_JOBS[job_id],
            "status": "done",
            "result": {
                "new": new_ones,
                "updated": updated_ones,
                "discharged": discharged_ones,
                "errors": errors,
                "parsed_count": len(parsed),
                "confirm": confirm,
            },
        }
    except Exception as e:
        logger.exception("Import job %s crashed", job_id)
        IMPORT_JOBS[job_id] = {
            **IMPORT_JOBS[job_id],
            "status": "failed",
            "error": f"Error inesperado: {str(e)[:300]}",
        }


async def _upsert_daily_from_censo(patient_id: str, user_id: str, target_date: str, entry: dict):
    """Merge labs/studies/procedures/cultures/vital_signs/indications from a census row.
    Non-destructive: appends without duplicating. Diffs indications against previous stored text."""
    labs = entry.get("labs")
    studies = entry.get("studies")
    procedures = entry.get("procedures")
    cultures = entry.get("cultures")
    vitals = entry.get("vital_signs")
    events = entry.get("events")
    new_ind = entry.get("_indications_text", "")
    if not any([labs, studies, procedures, cultures, vitals, events, new_ind]):
        return

    # Compute indications changes vs previously stored
    ind_changes = ""
    if new_ind:
        # Find most recent daily_entry with any indications_changes stored (or any daily) as reference
        prev_cur = db.daily_entries.find(
            {"patient_id": patient_id, "user_id": user_id, "date": {"$lt": target_date}},
            {"_id": 0},
        )
        prev_list = await prev_cur.to_list(30)
        prev_list.sort(key=lambda x: x.get("date", ""), reverse=True)
        prev_ind = ""
        # Store new_ind as raw on patient for future comparison
        # Compute simple line-diff: lines in new not in prev
        prev_patient = await db.patients.find_one(
            {"id": patient_id, "user_id": user_id}, {"_id": 0}
        )
        if prev_patient:
            prev_ind = (prev_patient.get("last_indications_text") or "").strip()
        new_lines = [ln.strip() for ln in new_ind.split("\n") if ln.strip()]
        prev_set = {ln.strip().lower() for ln in prev_ind.split("\n") if ln.strip()}
        added = [ln for ln in new_lines if ln.lower() not in prev_set]
        removed = [ln for ln in prev_ind.split("\n") if ln.strip() and ln.strip().lower() not in {n.lower() for n in new_lines}]
        chunks = []
        if added:
            chunks.append("Agregado / modificado:\n- " + "\n- ".join(added))
        if removed:
            chunks.append("Suspendido / retirado:\n- " + "\n- ".join(removed))
        ind_changes = "\n\n".join(chunks) if (added or removed) else ""
        # Store new indications baseline on the patient
        await db.patients.update_one(
            {"id": patient_id, "user_id": user_id},
            {"$set": {"last_indications_text": new_ind, "updated_at": now_utc().isoformat()}},
        )

    existing = await db.daily_entries.find_one(
        {"patient_id": patient_id, "user_id": user_id, "date": target_date}
    )
    now_iso = now_utc().isoformat()
    if not existing:
        de = DailyEntry(
            patient_id=patient_id, user_id=user_id, date=target_date,
            labs=labs or "", studies=studies or "",
            procedures=procedures or "", cultures=cultures or "",
            vital_signs=vitals or "",
            indications_changes=ind_changes or "",
            events=events or "",
        )
        await db.daily_entries.insert_one(de.model_dump())
        return
    update = {"updated_at": now_iso}
    for field, val in (
        ("labs", labs), ("studies", studies), ("procedures", procedures),
        ("cultures", cultures), ("vital_signs", vitals),
        ("indications_changes", ind_changes), ("events", events),
    ):
        if not val:
            continue
        existing_val = (existing.get(field) or "").strip()
        if existing_val and val.strip() not in existing_val:
            update[field] = existing_val + "\n" + val.strip()
        elif not existing_val:
            update[field] = val.strip()
    if len(update) > 1:
        await db.daily_entries.update_one(
            {"patient_id": patient_id, "user_id": user_id, "date": target_date},
            {"$set": update},
        )


# ---------------- Paciente sin cambios ----------------
@api.post("/patients/{patient_id}/generate/no-changes")
async def generate_no_changes(patient_id: str, body: GenerateBody, user: dict = Depends(get_user)):
    patient = await db.patients.find_one({"id": patient_id, "user_id": user["id"]}, {"_id": 0})
    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    target = body.date or today_iso()

    today_entry = await db.daily_entries.find_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": target}, {"_id": 0}
    ) or {"labs": "", "studies": "", "events": "", "dictation": ""}

    # Previous pase (most recent with ai_pase_summary before today)
    prev_cur = db.daily_entries.find(
        {"patient_id": patient_id, "user_id": user["id"], "date": {"$lt": target}, "ai_pase_summary": {"$ne": None}},
        {"_id": 0},
    )
    prev_list = await prev_cur.to_list(60)
    prev_list.sort(key=lambda x: x.get("date", ""), reverse=True)
    prev_pase = prev_list[0] if prev_list else None

    # Previous labs for comparison
    prev_labs_cur = db.daily_entries.find(
        {"patient_id": patient_id, "user_id": user["id"], "date": {"$lt": target}, "labs": {"$ne": ""}},
        {"_id": 0},
    )
    prev_labs_list = await prev_labs_cur.to_list(60)
    prev_labs_list.sort(key=lambda x: x.get("date", ""), reverse=True)
    prev_labs_entry = prev_labs_list[0] if prev_labs_list else None

    ctx = build_patient_context(patient)
    prev_pase_block = (
        f"\n\nPASE DEL DÍA PREVIO ({prev_pase['date']}):\n{prev_pase['ai_pase_summary']}"
        if prev_pase else "\n\n(No hay pase previo registrado.)"
    )
    prev_labs_block = (
        f"\n\nLABORATORIOS DEL DÍA PREVIO ({prev_labs_entry['date']}):\n{prev_labs_entry['labs']}"
        if prev_labs_entry else ""
    )

    # 1) PASE SIN CAMBIOS
    pase_prompt = f"""Este paciente NO tuvo eventualidades relevantes en las últimas 24 horas. NO se dictó nueva información porque el residente confirmó que continúa estable.

INFORMACIÓN FIJA DEL PACIENTE:
{ctx}

LABORATORIOS DE HOY ({target}): {today_entry.get('labs', '') or 'Sin nuevos laboratorios el día de hoy.'}
ESTUDIOS DE HOY: {today_entry.get('studies', '') or 'Ninguno.'}
{prev_labs_block}
{prev_pase_block}

Genera un RESUMEN DEL PASE (texto plano, sin markdown) con la MISMA estructura que se usa habitualmente:

▎ {patient.get('name','')}  |  Cama {patient.get('bed','ND')} · Piso {patient.get('floor','ND')}
▎ DEIH: {days_between(patient.get('admission_date'))} · DPQX: {days_between(patient.get('surgery_date'))}

DX: (dx resumido)
QX: (procedimiento y fecha)
APP RELEVANTES: (una línea; si no hay, omite)
ONCOLÓGICO: (si aplica; si no, omite)
INTERCONSULTAS ACTIVAS: (si hay; si no, omite)

RESUMEN CLÍNICO DEL DÍA:
- Paciente sin eventualidades durante las últimas 24 horas. (Adapta esta frase al contexto clínico específico: p. ej. "Paciente estable postquirúrgico, sin nuevos eventos.", "Continúa asintomática, sin cambios en el manejo.", etc.)
- (Menciona si continúa con esquema previo, tolerancia a dieta, control del dolor, herida, drenajes, según lo esperable si no hubo eventos.)

CAMBIOS DE MEDICAMENTOS:
- Sin cambios en el esquema.

ESTUDIOS / PROCEDIMIENTOS:
- {'Ver estudios cargados' if today_entry.get('studies') else 'Ninguno el día de hoy.'}

CAMBIOS RELEVANTES EN LABORATORIOS:
- (Si hay labs de hoy y previos, identifica SOLO cambios clínicamente relevantes con tendencias. Si NO hay labs de hoy, escribe: "Sin nuevos laboratorios el día de hoy.")

⚡ SUGERENCIA DE IA — PLAN SUGERIDO:
- (2-4 bullets con propuestas de continuidad. Ej: "Continuar mismo plan terapéutico", "Vigilar tolerancia a la vía oral", "Progresar según protocolo".)
— Requiere validación del médico responsable —

PENDIENTES / A DISCUTIR:
- (Reutiliza pendientes del pase previo si aplican; si ya se resolvieron, indica: "Sin pendientes activos.")

Reglas estrictas:
1. NO inventes eventos, complicaciones, signos ni síntomas que no estén en el pase previo o en los nuevos labs.
2. La consigna es "continúa igual"; refleja estabilidad clínica.
3. Español médico profesional. Sin adornos."""

    system_pase = "Eres un asistente experto para un residente de cirugía general. Redactas pases para pacientes ESTABLES sin eventualidades, reutilizando el contexto previo. Nunca inventas eventos ni complicaciones."
    pase_text = await llm_generate(system_pase, pase_prompt)

    # 2) NOTA MEDSYS — Prompt 4: few-shots permanentes del servicio + user examples
    is_uti = _is_uti_patient(patient)
    training = await _get_training_examples(user["id"])
    user_style_examples = training.get("uti_notes_examples" if is_uti else "hospital_notes_examples", "")
    user_style_block = _style_block(user_style_examples, "NOTAS DEL PROPIO USUARIO")
    builtin_style_block = uti_note_style_block() if is_uti else hospitalization_note_style_block()
    plantilla = "UTI/UTIM" if is_uti else "HOSPITALIZACIÓN (PISO)"

    if is_uti:
        format_rules = (
            "FORMATO OBLIGATORIO (imitando los ejemplos permanentes de UTI/UTIM del servicio):\n"
            "1. Inicia con 'NOTA DE EVOLUCIÓN POR CIRUGÍA GENERAL EN UTI' o 'EN UTIM'.\n"
            "2. Segundo párrafo 'SE TRATA DE ...' (edad, nombre, dx, fecha QX, hallazgos).\n"
            "3. Tercer párrafo comienza con 'AL PASE DE VISITA SE ENCUENTRA EN LO NEUROLÓGICO...' "
            "y enumera todos los sistemas EN ESTE ORDEN: NEUROLÓGICO → RESPIRATORIO → CARDIOVASCULAR → "
            "GASTRO METABÓLICO → HÍDRICO URINARIO → HEMATOINFECCIOSO. Todo en un solo párrafo corrido, "
            "reflejando estabilidad y continuidad de manejo.\n"
            "4. 'A LA EXPLORACIÓN FÍSICA ...' en párrafo corrido.\n"
            "5. 'PLAN: ...' párrafo corrido (rutina de UTI, vigilancia hemodinámica, cuantificación, "
            "toma de labs AM, continuar mismo esquema).\n"
            "6. 'PRONÓSTICO RESERVADO A EVOLUCIÓN' en línea sola.\n"
            "7. TODA la nota en MAYÚSCULAS. Sin listas, sin viñetas, sin numeración."
        )
    else:
        format_rules = (
            "FORMATO OBLIGATORIO (imitando los ejemplos permanentes de HOSPITALIZACIÓN del servicio):\n"
            "1. Inicia con 'NOTA DE EVOLUCIÓN POR CIRUGÍA GENERAL'.\n"
            "2. Segundo párrafo 'SE TRATA DE PACIENTE MASCULINO/FEMENINO DE X AÑOS, NOMBRE, "
            "CON DIAGNÓSTICO DE...'.\n"
            "3. Tercer párrafo comienza con 'AL PASE DE VISITA SE ENCUENTRA CON SIGNOS VITALES "
            "ESTABLES, DENTRO DE PARÁMETROS NORMALES...' e integra en párrafo corrido: dieta, dolor "
            "controlado, diuresis, evacuaciones, drenajes, antibióticos si aplica, reflejando "
            "estabilidad y ausencia de eventos en 24h.\n"
            "4. 'A LA EXPLORACIÓN FÍSICA ALERTA, REACTIVO, COOPERADOR...' en párrafo corrido.\n"
            "5. 'PLAN: ...' párrafo corrido (continuar mismo esquema, dieta, analgesia IV, "
            "cuidados de herida, cuantificación, vigilancia clínica).\n"
            "6. 'PRONÓSTICO RESERVADO A EVOLUCIÓN' en línea sola.\n"
            "7. TODA la nota en MAYÚSCULAS. Sin listas, sin viñetas, sin numeración."
        )

    note_prompt = (
        f"INFORMACIÓN FIJA DEL PACIENTE:\n{ctx}\n\n"
        f"RESUMEN DEL PASE DE HOY (RECIÉN GENERADO):\n{pase_text}\n\n"
        f"LABORATORIOS DE HOY: {today_entry.get('labs', '') or 'SIN NUEVOS LABORATORIOS.'}\n"
        f"ESTUDIOS DE HOY: {today_entry.get('studies', '') or 'NINGUNO.'}\n"
        f"PROCEDIMIENTOS: {today_entry.get('procedures', '') or 'NINGUNO.'}\n"
        f"CULTIVOS: {today_entry.get('cultures', '') or 'NINGUNO.'}\n"
        f"SIGNOS VITALES: {today_entry.get('vital_signs', '') or 'ND.'}\n\n"
        f"CONTEXTO: PACIENTE ESTABLE, SIN EVENTUALIDADES EN 24 HORAS.\n"
        f"TIPO DE NOTA A REDACTAR: {plantilla}\n\n"
        f"{format_rules}\n"
        f"{builtin_style_block}"
        f"{user_style_block}"
        f"\nAHORA REDACTA LA NOTA DEL PACIENTE INDICADO ARRIBA (paciente estable, sin cambios). "
        f"NUNCA copies texto de los ejemplos, solo IMITA el estilo. Devuelve SOLO la nota."
    )
    system_note = (
        "Eres RESIDENTE R3 DE CIRUGÍA GENERAL del HOSPITAL ÁNGELES redactando una nota MedSys "
        "para un paciente estable sin cambios. Debes escribir EXACTAMENTE con el estilo del "
        "servicio mostrado en los ejemplos permanentes.\n\n"
        "REGLAS INAMOVIBLES:\n"
        "- TODO EL TEXTO EN MAYÚSCULAS (excepto números y símbolos).\n"
        "- SIN LISTAS. SIN VIÑETAS. SIN NUMERACIÓN. Párrafos corridos separados por saltos de línea.\n"
        "- SIN MARKDOWN. Sin frases de IA. Sin 'basado en'. Sin 'ND' innecesarios.\n"
        "- Refleja estabilidad clínica. Nunca inventar datos.\n"
        "- Los ejemplos son solo REFERENCIA DE ESTILO — nunca copies nombres, diagnósticos, "
        "fechas ni frases textuales."
    )
    note_text = await llm_generate(system_note, note_prompt)

    # 3) WHATSAPP — Prompt 5: solo primer nombre + (cama), sin edad/dx/DPQX/DEIH
    wa_user_examples = _style_block(training.get("whatsapp_examples", ""), "MENSAJES DEL PROPIO USUARIO")
    wa_builtin_examples = whatsapp_style_block()
    full_name = (patient.get("name") or "").strip()
    first_name = full_name.split()[0].title() if full_name else "________"
    patient_bed = (patient.get("bed") or "").strip()
    wa_prompt = (
        f"Redacta un mensaje de WhatsApp corto (paciente estable, sin cambios) como lo enviaría un R3 "
        f"de Cirugía General del Hospital Ángeles al médico tratante. Un solo párrafo corrido, "
        f"capitalización normal (NO mayúsculas).\n\n"
        f"DATOS INTERNOS (usa solo lo necesario, NO copies nombre completo, edad, diagnóstico, "
        f"DPQX ni DEIH en el mensaje):\n{ctx}\n\n"
        f"Labs de hoy: {today_entry.get('labs', '') or 'sin nuevos labs.'}\n"
        f"Estudios: {today_entry.get('studies', '') or 'ninguno.'}\n\n"
        f"REGLAS OBLIGATORIAS:\n"
        f"- INICIA con: 'Buenos días Dr, pasé a ver a {first_name} ({patient_bed}), está con...' "
        f"o variante muy cercana.\n"
        f"- Solo primer nombre ({first_name}). NUNCA nombre completo, edad, dx, DPQX ni DEIH.\n"
        f"- Interpreta signos vitales: 'signos vitales estables', 'afebril', 'con tendencia a hipertensión'.\n"
        f"- Refleja estabilidad: sin cambios en 24h, tolera dieta, refiere buen control analgésico, "
        f"diuresis por micción espontánea, drenajes con gasto si aplica.\n"
        f"- Labs SOLO si cambian conducta. Nunca menciones labs normales.\n"
        f"- NUNCA uses: 'OJO', 'llama la atención', 'se sugiere', 'se recomienda', 'se informa que'.\n"
        f"- Si hay decisión pendiente termina con UNA sola pregunta: '¿Gusta que continuemos mismo "
        f"esquema?', '¿Gusta que valoremos egreso?', '¿Gusta que iniciemos vía oral?'. "
        f"Si no, cierre natural: 'Quedo al pendiente, saludos!', 'Lindo día doc!', 'Saludos y bonito día'.\n"
        f"- Un solo párrafo corrido. Sin listas. Sin markdown.\n"
        f"- Nunca inventes datos.\n"
        f"{wa_builtin_examples}"
        f"{wa_user_examples}"
        f"\nDevuelve SOLO el mensaje."
    )
    system_wa = (
        "Eres residente R3 de Cirugía General del Hospital Ángeles enviando WhatsApp corto al "
        "médico tratante para un paciente estable. Solo primer nombre y cama del paciente — "
        "NUNCA nombre completo, edad, diagnóstico, DPQX ni DEIH. Un solo párrafo corrido, "
        "capitalización normal. Interpreta signos vitales y labs. Nunca sonar a IA. "
        "Nunca 'OJO', 'llama la atención', 'se sugiere', 'se recomienda'."
    )
    wa_text = await llm_generate(system_wa, wa_prompt)

    # Save all
    await db.daily_entries.update_one(
        {"patient_id": patient_id, "user_id": user["id"], "date": target},
        {"$set": {
            "ai_pase_summary": pase_text,
            "ai_evolution_note": note_text,
            "ai_whatsapp": wa_text,
            "saved_to_pase": True,
            "updated_at": now_utc().isoformat(),
        }},
        upsert=True,
    )
    # First pase clears "nuevo ingreso" badge
    if patient.get("is_new_admission"):
        await db.patients.update_one(
            {"id": patient_id, "user_id": user["id"]},
            {"$set": {"is_new_admission": False, "updated_at": now_utc().isoformat()}},
        )
    return {"summary": pase_text, "note": note_text, "whatsapp": wa_text}


# ---------------- INGRESOS: nota de ingreso ----------------
@api.post("/patients/{patient_id}/generate/admission")
async def generate_admission_note(patient_id: str, body: GenerateBody, user: dict = Depends(get_user)):
    patient = await db.patients.find_one({"id": patient_id, "user_id": user["id"]}, {"_id": 0})
    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    ctx = build_patient_context(patient)
    target = body.date or today_iso()

    prompt = f"""Redacta una NOTA DE INGRESO hospitalario completa para este paciente, en español médico profesional, formato MedSys exacto listo para copiar. Usa exclusivamente la información fija almacenada.

INFORMACIÓN FIJA:
{ctx}

Formato (texto plano, sin markdown):

NOTA DE INGRESO — {target}
Servicio: Cirugía General

FICHA DE IDENTIFICACIÓN
Paciente: (nombre completo, edad, sexo)
Cama / Piso / Servicio: (cama · piso · servicio)
Fecha de ingreso: ({patient.get('admission_date', 'ND')})

MOTIVO DE INGRESO
(1-2 líneas basadas en el dx resumido y completo)

PADECIMIENTO ACTUAL
(3-6 líneas describiendo el cuadro que motiva el ingreso, basado en dx completo. NO inventes fechas ni síntomas específicos que no estén en la info fija.)

ANTECEDENTES DE IMPORTANCIA
- Personales patológicos: (según info fija; si no hay, "Sin antecedentes de importancia")
- Quirúrgicos: (si aplica cirugía previa/planeada)
- Alérgicos: ({patient.get('allergies', 'Ninguna referida')})
- Medicamentos crónicos: ({patient.get('important_medications', 'Ninguno referido')})
- Oncológico: ({patient.get('oncology_status', 'No aplica')} / {patient.get('oncology_treatment', '')})

EXPLORACIÓN FÍSICA AL INGRESO
Signos vitales: ND
Estado general: (impresión clínica al ingreso — general)
Cabeza y cuello: sin alteraciones referidas
Cardiopulmonar: sin alteraciones referidas
Abdomen: (según dx si aplica dolor, distensión, defensa)
Extremidades: sin alteraciones referidas
Neurológico: consciente, orientado, sin déficit referido

DIAGNÓSTICOS DE INGRESO
1. {patient.get('dx_short', 'ND')}
(Si hay dx_full, expandir en subincisos numerados con los diagnósticos secundarios relevantes.)

PLAN DE MANEJO INICIAL
1. Ingreso a servicio de Cirugía General.
2. (Ayuno / dieta según dx; ej. NPO si es abdomen agudo)
3. Soluciones parenterales / hidratación IV.
4. Analgesia según protocolo institucional.
5. Antibiótico si aplica según dx (mencionar si el dx sugiere infección).
6. Solicitar exámenes de laboratorio de ingreso (BH, QS, ES, TP/TTP, EGO, cultivos si aplica).
7. Estudios de imagen según dx (USG / TC abdominal si aplica).
8. Preparación quirúrgica si aplica.
9. Vigilancia clínica estrecha.
10. Interconsultas según requerimientos: ({patient.get('consultants', 'Ninguna al momento')})

⚡ SUGERENCIA DE IA — PLAN QUIRÚRGICO PROPUESTO:
- (1-3 bullets con la propuesta más razonable basada en dx; solo si el dx lo justifica.)
— Requiere validación del médico responsable —

PRONÓSTICO INICIAL
Reservado a evolución. (Ajusta según severidad del dx.)

—
Firma: Residente de Cirugía General

Reglas estrictas:
1. NO inventes signos vitales, fechas, medicamentos específicos ni valores.
2. Basa todo en la INFORMACIÓN FIJA proporcionada.
3. Si un dato no está, usa "ND" o "sin datos".
4. Todo lo terapéutico específico va en el PLAN o en la SUGERENCIA DE IA claramente etiquetada."""

    system = "Eres experto en redacción de notas de ingreso hospitalario de cirugía general para MedSys. No inventas datos clínicos. Español médico profesional."
    note = await llm_generate(system, prompt)
    return {"note": note}


# ---------------- NOTA DE INGRESO (pegada desde MedSys) ----------------
class AdmissionNotePaste(BaseModel):
    text: str


@api.post("/patients/{patient_id}/admission-note")
async def paste_admission_note(
    patient_id: str, body: AdmissionNotePaste, user: dict = Depends(get_user)
):
    """Guarda la nota de ingreso pegada desde MedSys y usa GPT para extraer campos fijos
    (merge no-destructivo — nunca borra información previa)."""
    patient = await db.patients.find_one({"id": patient_id, "user_id": user["id"]}, {"_id": 0})
    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    text = (body.text or "").strip()
    if not text or len(text) < 20:
        raise HTTPException(status_code=400, detail="El texto de la nota está vacío o es demasiado corto")

    # Always save the raw text first (permanent)
    await db.patients.update_one(
        {"id": patient_id, "user_id": user["id"]},
        {"$set": {"admission_note_text": text, "updated_at": now_utc().isoformat()}},
    )

    # Extract structured fields with GPT
    system = (
        "Eres un asistente que EXTRAE información clínica de una nota de ingreso hospitalario pegada literalmente. "
        "Devuelves EXCLUSIVAMENTE JSON válido. Nunca inventas datos. Si un campo no está, usas null."
    )
    prompt = f"""Extrae de esta NOTA DE INGRESO los siguientes campos.

NOTA DE INGRESO (texto tal como fue pegado):
{text}

Devuelve JSON con esta forma exacta:
{{
  "admission_reason": "Motivo de ingreso (una línea si está presente)",
  "current_illness": "Padecimiento actual completo tal como está escrito",
  "admission_diagnosis": "Diagnóstico(s) de ingreso principal(es)",
  "medical_history": "Antecedentes personales patológicos importantes (crónicos, degenerativos, etc.)",
  "previous_surgeries": "Cirugías previas si están mencionadas",
  "allergies": "Alergias específicas (medicamentos o sustancias)",
  "important_medications": "Medicamentos habituales / crónicos",
  "surgery_procedure": "Procedimiento quirúrgico realizado / planeado si está en la nota",
  "surgery_findings": "Hallazgos quirúrgicos si están mencionados",
  "oncology_status": "Estado oncológico si aplica (estadio, tipo, actividad de la enfermedad)"
}}

Reglas ESTRICTAS:
1. NO inventes nada. Si un campo NO está en la nota → null.
2. Copia el texto literal cuando corresponda; no reformules ni resumas en exceso.
3. NO extraigas ni menciones interconsultantes (esos los captura el usuario manualmente).
4. Devuelve SOLO el JSON, sin ```json``` ni explicaciones."""

    try:
        raw = await llm_generate(system, prompt)
    except Exception as e:
        return {
            "ok": True,
            "extracted": {},
            "merged_fields": [],
            "warning": f"Nota guardada, pero la extracción con IA falló: {str(e)[:200]}",
        }

    raw = raw.strip()
    if raw.startswith("```"):
        parts = raw.split("```", 2)
        raw = parts[1] if len(parts) > 1 else ""
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip("`").strip()

    import json as _json
    try:
        data = _json.loads(raw)
    except Exception:
        import re
        m = re.search(r"\{.*\}", raw, re.S)
        data = _json.loads(m.group(0)) if m else {}

    # Non-destructive merge: only fill fields that are currently empty/null
    FIELD_MAP = {
        "current_illness": "dx_full",
        "admission_diagnosis": "dx_short",
        "medical_history": "medical_history",
        "allergies": "allergies",
        "important_medications": "important_medications",
        "surgery_procedure": "surgery_procedure",
        "surgery_findings": "surgery_findings",
        "oncology_status": "oncology_status",
    }
    merged_fields: list[str] = []
    to_update: dict = {}
    for src, dst in FIELD_MAP.items():
        val = data.get(src)
        if not val or not isinstance(val, str) or not val.strip():
            continue
        current = (patient.get(dst) or "").strip()
        if not current:
            to_update[dst] = val.strip()
            merged_fields.append(dst)
        elif val.strip() != current and len(val.strip()) > len(current):
            # Append if new content adds substantially more info
            to_update[dst] = current + "\n\n[Nota de ingreso]: " + val.strip()
            merged_fields.append(dst)

    if to_update:
        to_update["updated_at"] = now_utc().isoformat()
        await db.patients.update_one(
            {"id": patient_id, "user_id": user["id"]}, {"$set": to_update}
        )

    return {"ok": True, "extracted": data, "merged_fields": merged_fields}


# ---------------- NOTAS ADICIONALES (otros servicios) ----------------
class AdditionalNoteCreate(BaseModel):
    source: str  # e.g. "Medicina Interna", "UTI", "Oncología"
    text: str


@api.get("/patients/{patient_id}/additional-notes")
async def list_additional_notes(patient_id: str, user: dict = Depends(get_user)):
    patient = await db.patients.find_one({"id": patient_id, "user_id": user["id"]}, {"_id": 0})
    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    cur = db.additional_notes.find(
        {"patient_id": patient_id, "user_id": user["id"]}, {"_id": 0}
    )
    notes = await cur.to_list(500)
    notes.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return notes


@api.post("/patients/{patient_id}/additional-notes")
async def add_additional_note(
    patient_id: str, body: AdditionalNoteCreate, user: dict = Depends(get_user)
):
    patient = await db.patients.find_one({"id": patient_id, "user_id": user["id"]}, {"_id": 0})
    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    text = (body.text or "").strip()
    source = (body.source or "Otro servicio").strip() or "Otro servicio"
    if not text or len(text) < 20:
        raise HTTPException(status_code=400, detail="El texto de la nota está vacío o es demasiado corto")

    # AI executive summary — does NOT modify patient's fixed information
    system = (
        "Eres un asistente que genera resúmenes ejecutivos de notas médicas de otros servicios. "
        "Solo resumes lo que aparece en la nota; NO inventas ni modificas información. "
        "Español médico profesional."
    )
    prompt = f"""Genera un RESUMEN EJECUTIVO breve de esta nota clínica proveniente de otro servicio médico ({source}).

NOTA COMPLETA:
{text}

Formato (texto plano, sin markdown), máximo 12 líneas:

RESUMEN — {source}
Impresión diagnóstica: (1-2 líneas)
Cambios importantes: (bullets con lo relevante clínicamente)
Recomendaciones: (bullets del servicio)
Pendientes: (bullets)

Reglas:
1. NO inventes. Solo lo que está en la nota.
2. Si algo no está, escribe "ND" o omite la línea.
3. Español médico. Sin comentarios extra."""

    ai_summary = None
    warning = None
    try:
        ai_summary = await llm_generate(system, prompt)
    except Exception as e:
        warning = f"Nota guardada, pero el resumen con IA falló: {str(e)[:200]}"

    note = {
        "id": str(uuid.uuid4()),
        "patient_id": patient_id,
        "user_id": user["id"],
        "source": source,
        "text": text,
        "ai_summary": ai_summary,
        "created_at": now_utc().isoformat(),
    }
    await db.additional_notes.insert_one(note)
    note.pop("_id", None)
    return {"note": note, "warning": warning}


@api.delete("/patients/{patient_id}/additional-notes/{note_id}")
async def delete_additional_note(
    patient_id: str, note_id: str, user: dict = Depends(get_user)
):
    result = await db.additional_notes.delete_one(
        {"id": note_id, "patient_id": patient_id, "user_id": user["id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Nota no encontrada")
    return {"ok": True}





def get_stt():
    return client


@api.post("/transcribe")
async def transcribe(
    audio: UploadFile = File(...),
    user: dict = Depends(get_user),
):
    """Real transcription..."""
    try:
        content = await audio.read()

        if len(content) < 1000:
            raise HTTPException(status_code=400, detail="Audio demasiado corto")

        filename = audio.filename or "audio.webm"

        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in ("mp3", "mp4", "mpeg", "mpga", "m4a", "wav", "webm"):
            filename = "audio.webm"

        audio_file = io.BytesIO(content)
        audio_file.name = filename

        transcript = await openai_client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            language="es",
            prompt="Transcripción médica en español. Terminología de cirugía general, laboratorios, medicamentos."
        )

        text = transcript.text
        return {"text": text.strip()}

    except HTTPException:
        raise
        
    except Exception as e:
        logger.exception("Transcription failed")
        raise HTTPException(
            status_code=500,
            detail=f"Error al transcribir: {str(e)}"
        )


@api.get("/")
async def root():
    return {"app": "SurgiNote", "status": "ok"}


async def seed_demo_user():
    existing = await db.users.find_one({"email": "demo@surginote.app"})
    if existing:
        return
    user_id = str(uuid.uuid4())
    await db.users.insert_one({
        "id": user_id,
        "email": "demo@surginote.app",
        "password": hash_pw("DemoSurgi2026!"),
        "name": "Dr. Demo",
        "created_at": now_utc().isoformat(),
    })
    logger.info("Seeded demo user")


app.include_router(api)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def on_start():
    await seed_demo_user()


@app.on_event("shutdown")
async def on_stop():
    client.close()
